import streamlit as st
from groq import Groq
import fitz
import base64
from PIL import Image
import io
import streamlit.components.v1 as components
import uuid
from datetime import datetime
import requests
from urllib.parse import urlencode
import json
import os
import hashlib
import bcrypt
import re

# ── Office file support ──
def read_excel_file(raw: bytes, filename: str) -> str:
    """Baca file Excel (.xlsx/.xls) dan convert ke teks terstruktur."""
    try:
        import openpyxl
        from io import BytesIO
        wb = openpyxl.load_workbook(BytesIO(raw), data_only=True)
        result = [f"[EXCEL: {filename} | {len(wb.sheetnames)} sheet]\n"]
        for sheet_name in wb.sheetnames[:5]:  # max 5 sheet
            ws = wb[sheet_name]
            result.append(f"\n=== Sheet: {sheet_name} ===")
            rows_added = 0
            for row in ws.iter_rows(values_only=True):
                if rows_added > 200:  # max 200 baris per sheet
                    result.append("...[data terpotong]")
                    break
                row_vals = [str(v) if v is not None else "" for v in row]
                if any(v.strip() for v in row_vals):  # skip baris kosong
                    result.append(" | ".join(row_vals))
                    rows_added += 1
        return "\n".join(result)[:40000]
    except Exception as e:
        return f"[Gagal baca Excel: {e}]"

def read_word_file(raw: bytes, filename: str) -> str:
    """Baca file Word (.docx) dan extract teks."""
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(raw))
        result = [f"[WORD: {filename}]\n"]
        for para in doc.paragraphs:
            if para.text.strip():
                result.append(para.text)
        # Baca tabel jika ada
        for i, table in enumerate(doc.tables[:10]):
            result.append(f"\n[Tabel {i+1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                result.append(" | ".join(cells))
        return "\n".join(result)[:40000]
    except Exception as e:
        return f"[Gagal baca Word: {e}]"

def create_excel_download(content_text: str, filename: str = "sigma_output.xlsx") -> bytes:
    """Buat file Excel dari teks analisa SIGMA."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from io import BytesIO
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Analisa SIGMA"
        # Header
        ws["A1"] = "SIGMA — Analisa Output"
        ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
        ws["A1"].fill = PatternFill("solid", fgColor="1B2A4A")
        ws.column_dimensions["A"].width = 20
        ws.column_dimensions["B"].width = 60
        # Isi konten baris per baris
        row = 3
        for line in content_text.split("\n"):
            line = line.strip()
            if not line:
                row += 1
                continue
            if line.startswith("---") or line.startswith("═"):
                row += 1
                continue
            if ":" in line:
                parts = line.split(":", 1)
                ws.cell(row=row, column=1, value=parts[0].strip()).font = Font(bold=True)
                ws.cell(row=row, column=2, value=parts[1].strip())
            else:
                ws.merge_cells(f"A{row}:B{row}")
                ws.cell(row=row, column=1, value=line)
            ws.cell(row=row, column=1).alignment = Alignment(wrap_text=True)
            ws.cell(row=row, column=2).alignment = Alignment(wrap_text=True)
            row += 1
        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except Exception:
        return b""

def create_word_download(content_text: str, filename: str = "sigma_output.docx") -> bytes:
    """Buat file Word dari teks analisa SIGMA."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        doc = Document()
        doc.add_heading("SIGMA — Analisa Output", level=0)
        for line in content_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("---") or line.startswith("═"):
                doc.add_paragraph("─" * 40)
            elif line.startswith(("📋","💰","🛡️","📊","💎","🏆","⚖️","📈","🔭")):
                p = doc.add_heading(line, level=2)
            elif ":" in line and len(line) < 80:
                parts = line.split(":", 1)
                p = doc.add_paragraph()
                run1 = p.add_run(parts[0].strip() + ": ")
                run1.bold = True
                p.add_run(parts[1].strip())
            else:
                doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return b""

# ─────────────────────────────────────────────
# MARKET CONTEXT — Real-time data injector
# ─────────────────────────────────────────────
def _fetch_stock_data(ticker: str) -> dict:
    """Ambil data OHLCV + info dasar via yfinance."""
    try:
        import yfinance as yf
        t = yf.Ticker(ticker)
        hist = t.history(period="5d")
        if hist.empty:
            return {}
        last = hist.iloc[-1]
        prev = hist.iloc[-2] if len(hist) > 1 else last
        chg = ((last["Close"] - prev["Close"]) / prev["Close"] * 100) if prev["Close"] else 0
        info = t.info
        return {
            "ticker": ticker,
            "close": round(last["Close"], 0),
            "volume": int(last["Volume"]),
            "change_pct": round(chg, 2),
            "pe_ratio": info.get("trailingPE", "N/A"),
            "market_cap": info.get("marketCap", "N/A"),
            "week52_high": info.get("fiftyTwoWeekHigh", "N/A"),
            "week52_low": info.get("fiftyTwoWeekLow", "N/A"),
        }
    except Exception:
        return {}

# Sumber berita finansial Indonesia — RSS gratis
NEWS_SOURCES = {
    "google":  "https://news.google.com/rss/search?q={query}&hl=id&gl=ID&ceid=ID:id",
    "cnbc_id": "https://www.cnbcindonesia.com/rss",
    "bisnis":  "https://ekonomi.bisnis.com/rss",
    "kontan":  "https://rss.kontan.co.id/category/investasi",
    "idx_news":"https://www.idx.co.id/umbraco/Surface/RssHelper/GetRss?categoryId=1",
}

def _fetch_news_headlines(query: str, max_items: int = 5) -> list:
    """
    Ambil headline berita dari multi-sumber:
    1. Google News (query-based)
    2. CNBC Indonesia
    3. Bisnis.com
    4. Kontan
    5. IDX News
    Deduplikasi judul yang mirip, return max_items per sumber.
    """
    try:
        import feedparser
        all_headlines = []
        seen_titles = set()

        # 1. Google News — paling relevan karena query-based
        try:
            url = NEWS_SOURCES["google"].format(
                query=requests.utils.quote(f"{query} saham Indonesia")
            )
            feed = feedparser.parse(url)
            for entry in feed.entries[:max_items]:
                title = entry.title.strip()
                key = title[:40].lower()
                if key not in seen_titles:
                    seen_titles.add(key)
                    pub = entry.get("published", "")[:16]
                    all_headlines.append(f"[Google News] [{pub}] {title}")
        except Exception:
            pass

        # 2. CNBC Indonesia
        try:
            feed = feedparser.parse(NEWS_SOURCES["cnbc_id"])
            count = 0
            for entry in feed.entries:
                if count >= 3:
                    break
                title = entry.title.strip()
                key = title[:40].lower()
                # Filter yang relevan dengan query
                if query.lower() in title.lower() or "saham" in title.lower() or "ihsg" in title.lower() or "bursa" in title.lower():
                    if key not in seen_titles:
                        seen_titles.add(key)
                        pub = entry.get("published", "")[:16]
                        all_headlines.append(f"[CNBC ID] [{pub}] {title}")
                        count += 1
        except Exception:
            pass

        # 3. Kontan
        try:
            feed = feedparser.parse(NEWS_SOURCES["kontan"])
            count = 0
            for entry in feed.entries:
                if count >= 3:
                    break
                title = entry.title.strip()
                key = title[:40].lower()
                if query.lower() in title.lower() or "saham" in title.lower() or "investasi" in title.lower():
                    if key not in seen_titles:
                        seen_titles.add(key)
                        pub = entry.get("published", "")[:16]
                        all_headlines.append(f"[Kontan] [{pub}] {title}")
                        count += 1
        except Exception:
            pass

        # 4. Bisnis.com
        try:
            feed = feedparser.parse(NEWS_SOURCES["bisnis"])
            count = 0
            for entry in feed.entries:
                if count >= 3:
                    break
                title = entry.title.strip()
                key = title[:40].lower()
                if query.lower() in title.lower() or "emiten" in title.lower() or "bursa" in title.lower():
                    if key not in seen_titles:
                        seen_titles.add(key)
                        pub = entry.get("published", "")[:16]
                        all_headlines.append(f"[Bisnis] [{pub}] {title}")
                        count += 1
        except Exception:
            pass

        return all_headlines[:max_items * 2]  # max 10 berita total

    except Exception:
        return []

def _fetch_market_news_general() -> list:
    """
    Ambil berita pasar modal umum (IHSG, makro, kebijakan)
    tanpa filter query — untuk konteks market overview.
    """
    try:
        import feedparser
        headlines = []
        seen = set()

        sources = [
            ("CNBC ID", NEWS_SOURCES["cnbc_id"]),
            ("Kontan",  NEWS_SOURCES["kontan"]),
            ("Bisnis",  NEWS_SOURCES["bisnis"]),
        ]

        market_keywords = ["ihsg", "bursa", "saham", "bi rate", "inflasi",
                           "rupiah", "idx", "ojk", "dividen", "emiten",
                           "right issue", "ipo", "buyback"]

        for source_name, url in sources:
            try:
                feed = feedparser.parse(url)
                count = 0
                for entry in feed.entries:
                    if count >= 3:
                        break
                    title = entry.title.strip()
                    key = title[:40].lower()
                    if any(kw in title.lower() for kw in market_keywords):
                        if key not in seen:
                            seen.add(key)
                            pub = entry.get("published", "")[:16]
                            headlines.append(f"[{source_name}] [{pub}] {title}")
                            count += 1
            except Exception:
                continue

        return headlines[:8]
    except Exception:
        return []

def _fetch_ihsg_data() -> dict:
    """Ambil data IHSG (^JKSE) sebagai konteks makro pasar."""
    try:
        import yfinance as yf
        t = yf.Ticker("^JKSE")
        hist = t.history(period="5d")
        if hist.empty:
            return {}
        last = hist.iloc[-1]
        prev = hist.iloc[-2] if len(hist) > 1 else last
        chg = ((last["Close"] - prev["Close"]) / prev["Close"] * 100) if prev["Close"] else 0
        return {
            "level": round(last["Close"], 0),
            "change_pct": round(chg, 2),
            "volume": int(last["Volume"]),
        }
    except Exception:
        return {}

# ─────────────────────────────────────────────
# FUNDAMENTAL DATA FETCHER — untuk melengkapi data PDF
# ─────────────────────────────────────────────
# Peta nama perusahaan ke ticker IDX
COMPANY_TICKER_MAP = {
    "bank central asia": "BBCA", "bca": "BBCA",
    "bank rakyat indonesia": "BBRI", "bri": "BBRI",
    "bank mandiri": "BMRI", "mandiri": "BMRI",
    "bank negara indonesia": "BBNI", "bni": "BBNI",
    "bank bsi": "BRIS", "bank syariah indonesia": "BRIS",
    "telkom": "TLKM", "astra": "ASII",
    "unilever": "UNVR", "indofood": "INDF",
    "goto": "GOTO", "gojek": "GOTO",
    "adaro": "ADRO", "antam": "ANTM",
}

def detect_ticker_from_pdf(pdf_text: str) -> str:
    """Deteksi ticker saham dari teks PDF laporan keuangan."""
    import re
    text_lower = pdf_text.lower()[:3000]
    # Cek peta nama perusahaan
    for name, ticker in COMPANY_TICKER_MAP.items():
        if name in text_lower:
            return ticker
    # Cek format kode saham langsung (4 huruf kapital)
    matches = re.findall(r'([A-Z]{4})', pdf_text[:2000])
    for m in matches:
        if m not in {"PADA","YANG","ATAU","DARI","BANK","TBKK","ANAK","ASET"}:
            return m
    return ""

def fetch_fundamental_supplement(ticker: str) -> dict:
    """
    Ambil data fundamental pelengkap dari yfinance.
    Digunakan untuk melengkapi data yang tidak ada di PDF.
    """
    try:
        import yfinance as yf
        t = yf.Ticker(f"{ticker}.JK")
        info = t.info
        hist = t.history(period="5d")
        price = round(hist.iloc[-1]["Close"], 0) if not hist.empty else None

        # Ambil data historis untuk hitung tren
        hist_annual = t.history(period="5y", interval="1mo")

        result = {
            "ticker": ticker,
            "price": price,
            "pe_ratio": info.get("trailingPE"),
            "pb_ratio": info.get("priceToBook"),
            "market_cap": info.get("marketCap"),
            "shares_outstanding": info.get("sharesOutstanding"),
            "book_value": info.get("bookValue"),
            "dividend_yield": info.get("dividendYield"),
            "trailing_eps": info.get("trailingEps"),
            "52w_high": info.get("fiftyTwoWeekHigh"),
            "52w_low": info.get("fiftyTwoWeekLow"),
            "sector": info.get("sector", ""),
            "industry": info.get("industry", ""),
        }
        return {k: v for k, v in result.items() if v is not None}
    except Exception:
        return {}

def build_fundamental_context(pdf_text: str) -> str:
    """
    Deteksi emiten dari PDF, fetch data pelengkap dari yfinance,
    dan return string konteks untuk diinject ke prompt analisa.
    """
    ticker = detect_ticker_from_pdf(pdf_text)
    if not ticker:
        return ""

    supp = fetch_fundamental_supplement(ticker)
    if not supp:
        return ""

    today_str = datetime.now().strftime("%d %B %Y")
    lines = [f"\n=== DATA PASAR REAL-TIME — {ticker} ({today_str}) ==="]

    if supp.get("price"):
        lines.append(f"Harga Saham Terkini : Rp{supp['price']:,.0f}")
    if supp.get("pe_ratio"):
        lines.append(f"PER (P/E)           : {supp['pe_ratio']:.2f}×")
    if supp.get("pb_ratio"):
        lines.append(f"PBV (P/B)           : {supp['pb_ratio']:.2f}×")
    if supp.get("trailing_eps"):
        lines.append(f"EPS (TTM)           : Rp{supp['trailing_eps']:,.0f}")
    if supp.get("book_value"):
        lines.append(f"Book Value/Share    : Rp{supp['book_value']:,.0f}")
    if supp.get("market_cap"):
        mc_t = supp["market_cap"] / 1e12
        lines.append(f"Market Cap          : Rp{mc_t:.1f} triliun")
    if supp.get("dividend_yield"):
        lines.append(f"Dividend Yield      : {supp['dividend_yield']*100:.2f}%")
    if supp.get("52w_high") and supp.get("52w_low"):
        lines.append(f"52W High / Low      : Rp{supp['52w_high']:,.0f} / Rp{supp['52w_low']:,.0f}")

    lines.append(f"\nGunakan data di atas untuk melengkapi PER, PBV, harga wajar,")
    lines.append(f"dan valuasi yang tidak tersedia di laporan keuangan PDF.")
    lines.append("=== AKHIR DATA PASAR ===")

    return "\n".join(lines)


# ─────────────────────────────────────────────
# MULTI-SOURCE DATA FETCHER — sistem backup berlapis
# ─────────────────────────────────────────────

def _fetch_price_stooq(ticker: str) -> dict:
    """Layer 2: Ambil harga historis dari stooq.com sebagai backup yfinance."""
    try:
        import pandas_datareader as pdr
        from datetime import timedelta
        end = datetime.now()
        start = end - timedelta(days=365*5)  # 5 tahun historis
        df = pdr.get_data_stooq(f"{ticker}.JK", start=start, end=end)
        if df.empty:
            return {}
        df = df.sort_index()
        last = df.iloc[-1]
        prev = df.iloc[-2] if len(df) > 1 else last
        chg = ((last["Close"] - prev["Close"]) / prev["Close"] * 100) if prev["Close"] else 0
        # Hitung simple metrics dari price history
        return {
            "price": round(last["Close"], 0),
            "change_pct": round(chg, 2),
            "52w_high": round(df["High"].tail(252).max(), 0),
            "52w_low": round(df["Low"].tail(252).min(), 0),
            "source": "stooq"
        }
    except Exception:
        return {}

def _fetch_idx_company_info(ticker: str) -> dict:
    """Layer 3: Ambil info dasar emiten dari IDX via request publik."""
    try:
        url = f"https://www.idx.co.id/umbraco/Surface/StockData/GetSecuritiesStock?start=0&length=1&keyword={ticker}"
        r = requests.get(url, timeout=5, headers={"User-Agent": "Mozilla/5.0"})
        if r.status_code != 200:
            return {}
        data = r.json()
        if not data.get("data"):
            return {}
        item = data["data"][0]
        return {
            "name": item.get("Name", ""),
            "sector": item.get("Sector", ""),
            "price": item.get("LastPrice", 0),
            "change_pct": item.get("Change", 0),
            "market_cap": item.get("MarketCap", 0),
            "source": "IDX"
        }
    except Exception:
        return {}

def _fetch_idx_financial_summary(ticker: str) -> str:
    """
    Ambil ringkasan keuangan terbaru dari IDX via API publik.
    Return string data atau kosong jika gagal.
    """
    try:
        # IDX Summary endpoint
        url = f"https://www.idx.co.id/umbraco/Surface/Helper/GetCompanyProfiles?stockCode={ticker}"
        r = requests.get(url, timeout=8, headers={
            "User-Agent": "Mozilla/5.0",
            "Referer": "https://www.idx.co.id/"
        })
        if r.status_code != 200:
            return ""
        data = r.json()
        if not data:
            return ""

        lines = [f"\n── DATA IDX TERBARU ({ticker}) ──"]
        if data.get("SektorEmiten"):
            lines.append(f"Sektor          : {data['SektorEmiten']}")
        if data.get("SubSektorEmiten"):
            lines.append(f"Sub-Sektor      : {data['SubSektorEmiten']}")
        if data.get("TanggalPencatatan"):
            lines.append(f"Tanggal IPO     : {data['TanggalPencatatan']}")
        if data.get("JumlahSaham"):
            shares = int(data["JumlahSaham"])
            lines.append(f"Jumlah Saham    : {shares/1e9:.2f} miliar lembar")

        return "\n".join(lines) if len(lines) > 1 else ""
    except Exception:
        return ""

def _fetch_idx_price_latest(ticker: str) -> dict:
    """Ambil harga terkini dari IDX API."""
    try:
        url = f"https://www.idx.co.id/umbraco/Surface/StockData/GetTradingInfoSS?code={ticker}"
        r = requests.get(url, timeout=5, headers={
            "User-Agent": "Mozilla/5.0",
            "Referer": "https://www.idx.co.id/"
        })
        if r.status_code != 200:
            return {}
        data = r.json()
        if not data:
            return {}
        return {
            "price": data.get("LastPrice", 0),
            "change": data.get("Change", 0),
            "change_pct": data.get("ChangePercentage", 0),
            "volume": data.get("Volume", 0),
            "market_cap": data.get("MarketCap", 0),
            "source": "IDX_live"
        }
    except Exception:
        return {}


def _fetch_price_with_fallback(ticker: str) -> dict:
    """
    Coba ambil data harga dengan urutan:
    1. IDX Live → 2. yfinance → 3. stooq → 4. IDX company API
    Return dict dengan field 'source' untuk tracking.
    """
    # Layer 0: IDX Live — paling fresh untuk harga
    try:
        idx_live = _fetch_idx_price_latest(ticker)
        if idx_live.get("price") and idx_live["price"] > 0:
            # Merge dengan yfinance untuk rasio (PER, PBV, dll)
            try:
                import yfinance as yf
                info = yf.Ticker(f"{ticker}.JK").info
                idx_live["pe_ratio"] = info.get("trailingPE")
                idx_live["pb_ratio"] = info.get("priceToBook")
                idx_live["eps"] = info.get("trailingEps")
                idx_live["book_value"] = info.get("bookValue")
                idx_live["dividend_yield"] = info.get("dividendYield")
                idx_live["shares"] = info.get("sharesOutstanding")
                idx_live["52w_high"] = info.get("fiftyTwoWeekHigh")
                idx_live["52w_low"] = info.get("fiftyTwoWeekLow")
                idx_live["source"] = "IDX_live+yfinance"
            except Exception:
                pass
            return {k: v for k, v in idx_live.items() if v is not None}
    except Exception:
        pass

    # Layer 1: yfinance
    try:
        import yfinance as yf
        t = yf.Ticker(f"{ticker}.JK")
        hist = t.history(period="5d", auto_adjust=True)
        info = t.info
        if not hist.empty:
            last = hist.iloc[-1]
            prev = hist.iloc[-2] if len(hist) > 1 else last
            chg = ((last["Close"] - prev["Close"]) / prev["Close"] * 100) if prev["Close"] else 0
            return {
                "price": round(last["Close"], 0),
                "change_pct": round(chg, 2),
                "pe_ratio": info.get("trailingPE"),
                "pb_ratio": info.get("priceToBook"),
                "market_cap": info.get("marketCap"),
                "eps": info.get("trailingEps"),
                "book_value": info.get("bookValue"),
                "dividend_yield": info.get("dividendYield"),
                "52w_high": info.get("fiftyTwoWeekHigh"),
                "52w_low": info.get("fiftyTwoWeekLow"),
                "shares": info.get("sharesOutstanding"),
                "source": "yfinance"
            }
    except Exception:
        pass

    # Layer 2: stooq
    stooq_data = _fetch_price_stooq(ticker)
    if stooq_data:
        return stooq_data

    # Layer 3: IDX API
    idx_data = _fetch_idx_company_info(ticker)
    if idx_data:
        return idx_data

    # Semua gagal
    return {"source": "failed"}

def _fetch_financials_with_fallback(ticker: str) -> dict:
    """
    Ambil laporan keuangan historis dengan fallback:
    1. yfinance income_stmt (API baru)
    2. yfinance financials (API lama)
    3. Return kosong — model akan hitung dari knowledge
    """
    result = {"inc": None, "bs": None, "cf": None, "source": "none"}
    try:
        import yfinance as yf
        t = yf.Ticker(f"{ticker}.JK")

        # Coba API baru dulu
        try:
            inc = t.income_stmt
            if inc is not None and not inc.empty:
                result["inc"] = inc
                result["source"] = "yfinance_new"
        except Exception:
            pass

        # Fallback ke API lama
        if result["inc"] is None:
            try:
                inc = t.financials
                if inc is not None and not inc.empty:
                    result["inc"] = inc
                    result["source"] = "yfinance_old"
            except Exception:
                pass

        # Balance sheet
        try:
            result["bs"] = t.balance_sheet
        except Exception:
            pass

        # Cashflow
        try:
            cf = t.cash_flow
            if cf is None or (hasattr(cf, "empty") and cf.empty):
                cf = t.cashflow
            result["cf"] = cf
        except Exception:
            pass

    except Exception:
        pass

    return result


def fetch_full_fundamental(ticker: str) -> str:
    """
    Tarik data fundamental lengkap dengan sistem backup berlapis:
    Layer 1: yfinance → Layer 2: stooq → Layer 3: IDX API → Layer 4: knowledge model
    """
    today_str = datetime.now().strftime("%d %B %Y")
    current_year = datetime.now().year

    # ── Ambil harga & rasio dengan fallback ──
    price_data = _fetch_price_with_fallback(ticker)
    price_source = price_data.get("source", "failed")

    # ── Ambil laporan keuangan dengan fallback ──
    fin_data = _fetch_financials_with_fallback(ticker)
    inc = fin_data.get("inc")
    bs  = fin_data.get("bs")
    cf  = fin_data.get("cf")
    fin_source = fin_data.get("source", "none")

    try:

        lines = []
        lines.append(f"=== DATA FUNDAMENTAL LENGKAP — {ticker} ({today_str}) ===")
        lines.append(f"Sumber harga : {price_source} | Sumber lapkeu: {fin_source}\n")

        # Tambahkan info IDX jika tersedia
        idx_info = _fetch_idx_financial_summary(ticker)
        if idx_info:
            lines.append(idx_info)
            lines.append("")

        # ── HARGA & VALUASI (dari price_data berlapis) ──
        lines.append("── HARGA & VALUASI TERKINI ──")
        if price_data.get("price"):
            lines.append(f"Harga Saham     : Rp{price_data['price']:,.0f}")
        if price_data.get("change_pct") is not None:
            arah = "▲" if price_data["change_pct"] >= 0 else "▼"
            lines.append(f"Perubahan       : {arah}{abs(price_data['change_pct']):.2f}%")
        if price_data.get("pe_ratio"):
            lines.append(f"PER (Trailing)  : {price_data['pe_ratio']:.2f}×")
        if price_data.get("pb_ratio"):
            lines.append(f"PBV             : {price_data['pb_ratio']:.2f}×")
        if price_data.get("market_cap"):
            lines.append(f"Market Cap      : Rp{price_data['market_cap']/1e12:.1f} triliun")
        if price_data.get("52w_high") and price_data.get("52w_low"):
            lines.append(f"52W High/Low    : Rp{price_data['52w_high']:,.0f} / Rp{price_data['52w_low']:,.0f}")
        if price_data.get("book_value"):
            lines.append(f"Book Value/Sh   : Rp{price_data['book_value']:,.0f}")
        if price_data.get("eps"):
            lines.append(f"EPS (Trailing)  : Rp{price_data['eps']:,.0f}")
        if price_data.get("dividend_yield"):
            lines.append(f"Dividend Yield  : {price_data['dividend_yield']*100:.2f}%")
        if price_data.get("shares"):
            lines.append(f"Shares Outstand : {price_data['shares']/1e9:.2f} miliar lembar")

        # Hitung PER & PBV manual jika tidak ada dari source
        _price = price_data.get("price", 0)
        _eps   = price_data.get("eps", 0)
        _bv    = price_data.get("book_value", 0)
        if _price and _eps and not price_data.get("pe_ratio"):
            lines.append(f"PER (hitung)    : {_price/_eps:.2f}× (Rp{_price:,.0f} ÷ Rp{_eps:,.0f})")
        if _price and _bv and not price_data.get("pb_ratio"):
            lines.append(f"PBV (hitung)    : {_price/_bv:.2f}× (Rp{_price:,.0f} ÷ Rp{_bv:,.0f})")

        if price_source == "failed":
            lines.append("⚠️ Semua sumber harga gagal — gunakan knowledge internal untuk estimasi harga")

        # ── INCOME STATEMENT HISTORIS ──
        if inc is not None and hasattr(inc, "empty") and not inc.empty:
            lines.append("\n── INCOME STATEMENT HISTORIS (tahunan, dalam miliar Rp) ──")
            # Ambil baris penting
            key_rows = {
                "Total Revenue"         : "Pendapatan",
                "Gross Profit"          : "Laba Kotor",
                "Operating Income"      : "Laba Operasional",
                "Net Income"            : "Laba Bersih",
                "Basic EPS"             : "EPS",
            }
            for yf_key, label in key_rows.items():
                try:
                    if yf_key in inc.index:
                        row = inc.loc[yf_key].dropna()
                        vals = []
                        for col in sorted(row.index, reverse=True)[:5]:
                            yr = str(col)[:4]
                            v = row[col]
                            if abs(v) > 1e9:
                                vals.append(f"{yr}: Rp{v/1e12:.1f}T")
                            elif abs(v) > 1e6:
                                vals.append(f"{yr}: Rp{v/1e9:.1f}M")
                            else:
                                vals.append(f"{yr}: {v:.0f}")
                        lines.append(f"{label:20s}: {' | '.join(vals)}")
                except Exception:
                    pass

        # ── BALANCE SHEET HISTORIS ──
        if bs is not None and hasattr(bs, "empty") and not bs.empty:
            lines.append("\n── BALANCE SHEET HISTORIS (tahunan, dalam triliun Rp) ──")
            key_rows_bs = {
                "Total Assets"              : "Total Aset",
                "Total Liabilities Net Minority Interest": "Total Liabilitas",
                "Stockholders Equity"       : "Total Ekuitas",
                "Common Stock Equity"       : "Ekuitas",
            }
            for yf_key, label in key_rows_bs.items():
                try:
                    if yf_key in bs.index:
                        row = bs.loc[yf_key].dropna()
                        vals = []
                        for col in sorted(row.index, reverse=True)[:5]:
                            yr = str(col)[:4]
                            v = row[col]
                            vals.append(f"{yr}: Rp{v/1e12:.1f}T")
                        lines.append(f"{label:25s}: {' | '.join(vals)}")
                except Exception:
                    pass

        # ── CASHFLOW HISTORIS ──
        if cf is not None and hasattr(cf, "empty") and not cf.empty:
            lines.append("\n── CASHFLOW HISTORIS (tahunan, dalam triliun Rp) ──")
            key_rows_cf = {
                "Operating Cash Flow"   : "Cash Flow Operasi",
                "Free Cash Flow"        : "Free Cash Flow",
                "Capital Expenditure"   : "Capex",
            }
            for yf_key, label in key_rows_cf.items():
                try:
                    if yf_key in cf.index:
                        row = cf.loc[yf_key].dropna()
                        vals = []
                        for col in sorted(row.index, reverse=True)[:5]:
                            yr = str(col)[:4]
                            v = row[col]
                            vals.append(f"{yr}: Rp{v/1e12:.1f}T")
                        lines.append(f"{label:25s}: {' | '.join(vals)}")
                except Exception:
                    pass

        # ── BERITA TERBARU EMITEN ──
        news_items = _fetch_news_headlines(ticker, max_items=6)
        if news_items:
            lines.append("\n── BERITA TERBARU ──")
            lines.extend(news_items)

        lines.append(f"\n=== PENTING — BACA SEBELUM ANALISA ===")
        lines.append(f"TAHUN SAAT INI: {current_year}")
        lines.append(f"PERINGATAN DATA: Data historis dari yfinance untuk saham IDX")
        lines.append(f"sering hanya tersedia sampai 2022-2023. Ini BUKAN data terbaru.")
        lines.append(f"WAJIB lakukan hal berikut:")
        lines.append(f"1. Data historis yfinance (2020-2022) → gunakan sebagai TREN HISTORIS LAMA")
        lines.append(f"2. Untuk tahun {current_year-1} dan {current_year} → gunakan knowledge training kamu")
        lines.append(f"   tentang emiten ini untuk estimasi, dan SEBUTKAN bahwa ini estimasi")
        lines.append(f"3. Tren 3 tahun TERAKHIR = {current_year-2}, {current_year-1}, {current_year}")
        lines.append(f"   Isi dengan: data yfinance untuk tahun yang ada, estimasi untuk yang tidak ada")
        lines.append(f"4. Setiap metrik di baris TERPISAH — DILARANG digabung satu baris")
        lines.append(f"5. Proyeksi {current_year+1}-{current_year+2} berdasarkan tren CAGR historis")
        lines.append(f"6. Di bagian Verdict, sebutkan sumber data: mana yang real-time, mana estimasi")
        lines.append(f"=== AKHIR INSTRUKSI ===")

        return "\n".join(lines)

    except Exception as e:
        return (
            f"[Data fetch gagal: {str(e)}]\n"
            f"Gunakan knowledge internal kamu tentang {ticker} untuk analisa fundamental. "
            f"Sebutkan bahwa data diambil dari knowledge training, bukan real-time."
        )


def get_market_context(prompt: str) -> str:
    """
    Deteksi kode saham dari prompt (format: 4 huruf kapital atau diikuti .JK),
    fetch data real-time, dan return string konteks untuk diinject ke prompt.
    Jika tidak ada saham spesifik, hanya inject kondisi IHSG + berita umum.
    """
    import re
    today_str = datetime.now().strftime("%d %B %Y, %H:%M WIB")

    # Deteksi ticker — 4 huruf kapital, opsional dengan .JK
    tickers_raw = re.findall(r'\b([A-Z]{4})(?:\.JK)?\b', prompt)
    # Filter kata umum yang bukan ticker
    SKIP = {"BBRI","BBCA","BMRI","TLKM","ASII","GOTO","BRIS","UNVR","ICBP","INDF",
            "ANTM","PTBA","ADRO","EXCL","SMGR","INTP","KLBF","SIDO","MYOR","CPIN"}
    # Ambil semua yang terdeteksi, prioritaskan yang memang saham IDX populer
    tickers = list(dict.fromkeys(tickers_raw))  # deduplicate, preserve order

    context_parts = [f"📅 Tanggal Analisa: {today_str}"]

    # IHSG
    ihsg = _fetch_ihsg_data()
    if ihsg:
        arah = "▲" if ihsg.get("change_pct", 0) >= 0 else "▼"
        context_parts.append(
            f"\n📈 IHSG: {ihsg['level']:,.0f} {arah}{abs(ihsg['change_pct'])}% | "
            f"Vol: {ihsg['volume']:,}"
        )

    # Data per saham yang disebut
    for raw_ticker in tickers[:3]:  # max 3 saham per query
        jk_ticker = f"{raw_ticker}.JK"
        stock = _fetch_stock_data(jk_ticker)
        if stock:
            arah = "▲" if stock.get("change_pct", 0) >= 0 else "▼"
            context_parts.append(
                f"\n📊 {raw_ticker}: Rp{stock['close']:,.0f} {arah}{abs(stock['change_pct'])}% | "
                f"Vol: {stock['volume']:,} | P/E: {stock['pe_ratio']} | "
                f"52W H/L: {stock['week52_high']}/{stock['week52_low']}"
            )

    # Berita — multi-sumber
    if tickers:
        news_query = f"saham {tickers[0]} IDX Indonesia"
        headlines = _fetch_news_headlines(news_query, max_items=4)
        if headlines:
            context_parts.append(f"\n📰 Berita {tickers[0]} (multi-sumber):")
            context_parts.extend(headlines)
    else:
        # Berita pasar umum jika tidak ada ticker spesifik
        general_news = _fetch_market_news_general()
        if general_news:
            context_parts.append(f"\n📰 Berita Pasar Modal Indonesia (terbaru):")
            context_parts.extend(general_news)

    if len(context_parts) <= 1:
        return ""  # Tidak ada data yang berhasil di-fetch

    return "\n".join(context_parts)

# ─────────────────────────────────────────────
# VERSION INFO
# ─────────────────────────────────────────────
SIGMA_VERSION = "1.4.0"
SIGMA_CHANGELOG = {
    "1.0.0": "Initial release — chat + analisa chart",
    "1.1.0": "Tambah real-time market data (yfinance)",
    "1.2.0": "Tambah analisa fundamental PDF",
    "1.3.0": "Multi-source data: yfinance + stooq + IDX API",
    "1.4.0": "Multi-source news, Excel/Word support, versioning",
}

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="KIPM SIGMA",
    layout="wide",
    initial_sidebar_state="expanded"
)

DATA_DIR = ".sigma_data"
os.makedirs(DATA_DIR, exist_ok=True)

# ─────────────────────────────────────────────
# PERSISTENCE
# ─────────────────────────────────────────────
def _ukey(email): return hashlib.md5(email.encode()).hexdigest()

def save_user(email, data):
    try:
        with open(os.path.join(DATA_DIR, f"{_ukey(email)}.json"), "w") as f:
            json.dump(data, f, ensure_ascii=False)
    except: pass

def load_user(email):
    try:
        p = os.path.join(DATA_DIR, f"{_ukey(email)}.json")
        if os.path.exists(p):
            with open(p) as f: return json.load(f)
    except: pass
    return None

# Username/password auth
def get_accounts():
    p = os.path.join(DATA_DIR, "accounts.json")
    if os.path.exists(p):
        with open(p) as f: return json.load(f)
    return {}

def save_accounts(acc):
    with open(os.path.join(DATA_DIR, "accounts.json"), "w") as f:
        json.dump(acc, f)

def register_user(username, password, display_name):
    acc = get_accounts()
    if username in acc: return False, "Username sudah dipakai"
    hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
    acc[username] = {"password": hashed, "display_name": display_name, "email": f"{username}@local"}
    save_accounts(acc)
    return True, "Berhasil daftar"

def login_user(username, password):
    acc = get_accounts()
    if username not in acc: return None
    if bcrypt.checkpw(password.encode(), acc[username]["password"].encode()):
        return {"email": acc[username]["email"], "name": acc[username]["display_name"], "picture": ""}
    return None

# ─────────────────────────────────────────────
# THEME COLORS
# ─────────────────────────────────────────────
def get_colors(theme="dark"):
    dark = theme == "dark"
    return {
        "bg":           "#212121" if dark else "#f0f0f0",
        "sidebar_bg":   "#171717" if dark else "#e3e3e3",
        "text":         "#ececec" if dark else "#0d0d0d",
        "text_muted":   "#8e8ea0" if dark else "#6e6e80",
        "border":       "#2f2f2f" if dark else "#d0d0d0",
        "hover":        "#2f2f2f" if dark else "#d0d0d0",
        "input_bg":     "#2f2f2f" if dark else "#ffffff",
        "bubble":       "#1B2A4A",
        "bubble_text":  "#ffffff",
        "divider":      "#2f2f2f" if dark else "#d0d0d0",
        "gold":         "#F5C242",
        "active_bg":    "#2f2f2f" if dark else "#c8c8c8",
    }

# ─────────────────────────────────────────────
# SESSION INIT
# ─────────────────────────────────────────────
def init_session():
    defaults = {
        "user": None,
        "theme": "dark",
        "data_loaded": False,
        "sessions": None,
        "active_id": None,
        "rename_id": None,
        "img_data": None,
        "pdf_data": None,
        "show_settings": False,
        "auth_mode": "login",  # login | register | google
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session()

C = get_colors(st.session_state.theme)

# ─────────────────────────────────────────────
# SYSTEM PROMPT
# ─────────────────────────────────────────────
SYSTEM_PROMPT = {
    "role": "system",
    "content": """Kamu adalah SIGMA — asisten trading dan analis pasar modal dari KIPM Universitas Pancasila, dibuat oleh komunitas Market n Mocha (MnM).

KEPRIBADIAN:
- Saat ngobrol biasa: ramah, hangat, santai, natural — seperti teman trader yang asik diajak ngobrol
- Saat diminta analisa: berubah menjadi profesional, tajam, tegas, sangat berpengalaman
- Selalu Bahasa Indonesia yang natural, tidak kaku
- Jangan pernah memulai jawaban dengan menjejalkan data pasar kecuali diminta

CARA MERESPONS:
- Sapaan biasa → balas ramah, perkenalkan diri singkat, tanya ada yang bisa dibantu
- Pertanyaan umum saham/pasar → jawab informatif tapi conversational
- Ada gambar/chart → LANGSUNG analisa teknikal, tidak perlu nunggu perintah
- Ada PDF laporan keuangan → LANGSUNG analisa fundamental, deteksi sektor otomatis
- Ada perintah analisa teknikal → gunakan FORMAT TRADE PLAN
- Ada perintah analisa fundamental → gunakan FORMAT ANALISA FUNDAMENTAL

═══════════════════════════════════════
FRAMEWORK TEKNIKAL (MnM Strategy+):
═══════════════════════════════════════
1. IFVG — Inversion Fair Value Gap
2. FVG — Fair Value Gap
3. Order Block (OB)
4. Supply & Demand Zones
5. Moving Average (EMA 13/21/50)
6. Bandarmologi — akumulasi/distribusi, delta volume, anomali volume
7. Volume Profile — VPOC, VAH, VAL

FORMAT TRADE PLAN:
📊 TRADE PLAN — [SAHAM] ([TIMEFRAME])
⚡ Bias: [Bullish/Bearish/Sideways]
🎯 Entry: [harga]
🛑 Stop Loss: [harga]
✅ Target 1: [harga]
✅ Target 2: [harga]
📦 Bandarmologi: [ringkasan volume & aksi bandar]
⚠️ Invalidasi: [kondisi yang membatalkan setup]
⚠️ DYOR — bukan rekomendasi investasi

═══════════════════════════════════════
FRAMEWORK FUNDAMENTAL:
═══════════════════════════════════════
DETEKSI SEKTOR OTOMATIS:
- Jika dokumen mengandung: NPL, NIM, DPK, CAR, LDR, BOPO, kredit, dana pihak ketiga
  → gunakan FRAMEWORK PERBANKAN
- Jika tidak → gunakan FRAMEWORK UMUM

── FRAMEWORK UMUM (non-perbankan) ──
• Warren Buffett: ROE >15% konsisten 5-10 tahun, DER <0.5, FCF > Net Income, ada moat
• Benjamin Graham: PBV <1.5, PER <15, PER×PBV <22.5, EPS positif 5 tahun berturut
• Peter Lynch: PEG Ratio <1 ideal (<2 acceptable), revenue growth >20% YoY
• CAN SLIM: EPS quarter naik >25% YoY, annual EPS naik >25% 3 tahun berturut

── FRAMEWORK PERBANKAN ──
• NIM (Net Interest Margin): sehat >4%
• NPL (Non Performing Loan): sehat <3% — KRITIS jika >5%
• LDR (Loan to Deposit Ratio): ideal 80-92%
• CAR (Capital Adequacy Ratio): aman >14%, minimum BI 8%
• ROA: sehat >1.5%
• ROE: sehat >15%
• BOPO: efisien <70% — semakin kecil semakin baik
• CIR (Cost to Income Ratio): ideal <45%
• Tren DPK, kredit, dan laba bersih YoY
• Tren dividen dan payout ratio

FORMAT ANALISA FUNDAMENTAL:
Ikuti PERSIS format di bawah ini. Setiap baris adalah satu item. DILARANG menggabungkan beberapa item dalam satu baris.

---
📋 ANALISA FUNDAMENTAL — [NAMA EMITEN] ([PERIODE])
🏦 Sektor    : [Perbankan / Non-Perbankan]
📌 Framework : [Perbankan / Buffett / Graham]

---
💰 PROFITABILITAS

ROE         : [X]%  →  standar >15%      [✅/⚠️/❌]
ROA         : [X]%  →  standar >1,5%     [✅/⚠️/❌]
NIM         : [X]%  →  standar >4%       [✅/⚠️/❌]
Laba Bersih : Rp[X] triliun  →  YoY [+/-X]%
EPS         : Rp[X]  →  YoY [+/-X]%

---
🛡️ KUALITAS ASET & RISIKO

NPL Gross   : [X]%  →  sehat <3%         [✅/⚠️/❌]
NPL Net     : [X]%  →  sehat <1%         [✅/⚠️/❌]
CAR         : [X]%  →  aman >14%         [✅/⚠️/❌]
BOPO        : [X]%  →  efisien <70%      [✅/⚠️/❌]
CIR         : [X]%  →  ideal <45%        [✅/⚠️/❌]
LDR         : [X]%  →  ideal 80-92%      [✅/⚠️/❌]

---
📊 PERTUMBUHAN YoY

Total Aset  : Rp[X] triliun  →  [+/-X]%
Kredit      : Rp[X] triliun  →  [+/-X]%
DPK         : Rp[X] triliun  →  [+/-X]%
Laba Bersih : Rp[X] triliun  →  [+/-X]%
EPS         : Rp[X]  →  [+/-X]%

---
💎 VALUASI

PER         : [X]×  →  Graham <15        [✅/⚠️/❌]
PBV         : [X]×  →  Graham <1,5       [✅/⚠️/❌]
Harga Wajar : Rp[X] – Rp[X]

---
🏆 DIVIDEN

Total Dividen  : Rp[X] triliun
Payout Ratio   : [X]%
Konsistensi    : [naik / stabil / turun]

---
📈 TREN 3 TAHUN TERAKHIR
(isi dengan angka AKTUAL dari data, bukan placeholder)

Laba Bersih  : Rp[angka aktual tahun-2]T → Rp[angka aktual tahun-1]T → Rp[angka aktual tahun ini]T (CAGR ~X%)
EPS          : Rp[angka aktual tahun-2] → Rp[angka aktual tahun-1] → Rp[angka aktual tahun ini]
ROE          : [angka aktual]% → [angka aktual]% → [angka aktual]%
Total Aset   : Rp[angka aktual]T → Rp[angka aktual]T → Rp[angka aktual]T

---
🔭 PROYEKSI 1-2 TAHUN KE DEPAN
(hitung dari tren CAGR data historis di atas)

Asumsi Growth: ~X% per tahun (berdasarkan CAGR laba bersih)
EPS Est.     : [tahun+1]: Rp[hasil hitung] | [tahun+2]: Rp[hasil hitung]
Laba Est.    : [tahun+1]: Rp[hasil hitung]T | [tahun+2]: Rp[hasil hitung]T
Target Harga : Konservatif Rp[PER rendah × EPS] | Moderat Rp[PER median × EPS] | Optimis Rp[PER tinggi × EPS]
Basis Hitung : PER historis rata-rata × EPS proyeksi

---
⚖️ VERDICT

Skor Fundamental : [X]/10

Kekuatan :
→ [poin kekuatan 1]
→ [poin kekuatan 2]
→ [poin kekuatan 3]

Risiko :
→ [poin risiko 1]
→ [poin risiko 2]

Valuasi Saat Ini : [Undervalue / Fairvalue / Overvalue]

Kesimpulan :
[3-4 kalimat: kondisi bisnis, kualitas fundamental, posisi valuasi,
saran akumulasi atau wait & see. Bahasa jelas dan mudah dipahami.]

---
⚠️ DYOR — Analisa ini bukan rekomendasi investasi.
Proyeksi bersifat estimasi berdasarkan tren historis, bukan jaminan.
Keputusan investasi sepenuhnya tanggung jawab investor.

ATURAN KERAS OUTPUT — WAJIB DIIKUTI:
1. Setiap metrik HARUS di baris SENDIRI — DILARANG digabung dalam satu baris
2. SEMUA angka yang ada di dokumen WAJIB diisi — tidak boleh tulis "tidak tersedia" jika angka ada
3. Angka dalam jutaan Rupiah → konversi ke triliun (bagi 1.000.000), bulatkan 2 desimal
4. Hitung YoY secara manual: ((nilai baru - nilai lama) / nilai lama) x 100
5. Pisahkan setiap seksi dengan garis ---
6. Verdict ditulis lengkap 3-4 kalimat, informatif, tidak generik

HIERARKI SUMBER DATA (urutan prioritas):
1. Data dari PDF laporan keuangan → sumber utama
2. Data dari "=== DATA PASAR REAL-TIME ===" → gunakan untuk PER, PBV, harga saham, market cap
3. Hitung manual dari rumus jika kedua sumber di atas tidak lengkap:
   - PER  = Harga Saham ÷ EPS
   - PBV  = Harga Saham ÷ Book Value per Saham
   - Book Value/Share = Total Ekuitas ÷ Jumlah Saham Beredar
   - Payout Ratio = Total Dividen ÷ Laba Bersih × 100
   - DPS = Total Dividen ÷ Jumlah Saham Beredar
   - NIM = Pendapatan Bunga Bersih ÷ Rata-rata Aset Produktif × 100
   - ROA = Laba Sebelum Pajak ÷ Rata-rata Total Aset × 100
   - ROE = Laba Bersih ÷ Rata-rata Ekuitas × 100
4. Jika benar-benar tidak bisa dihitung → tulis "Tidak dapat dihitung dari data tersedia"
   DILARANG menulis hanya "Tidak tersedia" tanpa mencoba hitung dari rumus

═══════════════════════════════════════
ATURAN KERAS:
═══════════════════════════════════════
- Jangan injeksi data pasar ke percakapan biasa/sapaan
- Gambar masuk → analisa teknikal langsung
- PDF masuk → analisa fundamental langsung, deteksi sektor otomatis
- DYOR disclaimer wajib di setiap output analisa"""
}

# ─────────────────────────────────────────────
# CHAT SESSION HELPERS
# ─────────────────────────────────────────────
def new_session():
    return {
        "id": str(uuid.uuid4())[:8],
        "title": "Obrolan Baru",
        "messages": [SYSTEM_PROMPT],
        "created": datetime.now().strftime("%d/%m %H:%M")
    }

def init_chat():
    if not st.session_state.sessions:
        s = new_session()
        st.session_state.sessions = [s]
        st.session_state.active_id = s["id"]
    else:
        for s in st.session_state.sessions:
            if not s["messages"] or s["messages"][0].get("role") != "system":
                s["messages"].insert(0, SYSTEM_PROMPT)
            else:
                s["messages"][0] = SYSTEM_PROMPT

def restore_images_from_messages():
    """Restore gambar dari message ke session_state agar tampil setelah refresh."""
    if not st.session_state.sessions:
        return
    for sesi in st.session_state.sessions:
        for i, msg in enumerate(sesi.get("messages", [])):
            if msg.get("role") == "user" and msg.get("img_b64"):
                key = f"thumb_{sesi['id']}_{i-1}"
                if key not in st.session_state:
                    st.session_state[key] = (msg["img_b64"], msg.get("img_mime", "image/jpeg"))

def get_active():
    for s in st.session_state.sessions:
        if s["id"] == st.session_state.active_id:
            return s
    return st.session_state.sessions[0]
    for s in st.session_state.sessions:
        if s["id"] == st.session_state.active_id:
            return s
    return st.session_state.sessions[0]

def delete_session(sid):
    st.session_state.sessions = [s for s in st.session_state.sessions if s["id"] != sid]
    if not st.session_state.sessions:
        ns = new_session()
        st.session_state.sessions = [ns]
    if st.session_state.active_id == sid:
        st.session_state.active_id = st.session_state.sessions[0]["id"]

# ─────────────────────────────────────────────
# GOOGLE OAUTH
# ─────────────────────────────────────────────
def google_auth_url():
    params = {
        "client_id": st.secrets.get("GOOGLE_CLIENT_ID", ""),
        "redirect_uri": st.secrets.get("GOOGLE_REDIRECT_URI", ""),
        "response_type": "code",
        "scope": "openid email profile",
        "access_type": "offline",
        "prompt": "select_account",
    }
    return "https://accounts.google.com/o/oauth2/v2/auth?" + urlencode(params)

def handle_oauth(code):
    r = requests.post("https://oauth2.googleapis.com/token", data={
        "code": code,
        "client_id": st.secrets.get("GOOGLE_CLIENT_ID", ""),
        "client_secret": st.secrets.get("GOOGLE_CLIENT_SECRET", ""),
        "redirect_uri": st.secrets.get("GOOGLE_REDIRECT_URI", ""),
        "grant_type": "authorization_code",
    })
    if r.status_code != 200: return None
    token = r.json().get("access_token", "")
    if not token: return None
    u = requests.get("https://www.googleapis.com/oauth2/v2/userinfo",
                     headers={"Authorization": f"Bearer {token}"})
    return u.json() if u.status_code == 200 else None

# ─────────────────────────────────────────────
# HANDLE OAUTH CALLBACK
# ─────────────────────────────────────────────
if "code" in st.query_params and st.session_state.user is None:
    info = handle_oauth(st.query_params["code"])
    if info:
        st.session_state.user = info
        saved = load_user(info["email"])
        if saved:
            st.session_state.theme = saved.get("theme", "dark")
            if saved.get("sessions"):
                st.session_state.sessions = saved["sessions"]
                st.session_state.active_id = saved.get("active_id")
        st.session_state.data_loaded = True
        # Buat token untuk auto-login — sama seperti username login
        token = str(uuid.uuid4()).replace("-","")
        with open(os.path.join(DATA_DIR, f"token_{token}.json"), "w") as f:
            json.dump(info, f)
        st.session_state.current_token = token
        st.query_params.clear()
        st.query_params["sigma_token"] = token  # Set di URL agar persist saat refresh
        st.rerun()

# ─────────────────────────────────────────────
# RESTORE SESSION DARI TOKEN (auto-login saat refresh)
# ─────────────────────────────────────────────
if "sigma_token" in st.query_params and st.session_state.user is None:
    token = st.query_params.get("sigma_token", "")
    token_file = os.path.join(DATA_DIR, f"token_{token}.json")
    if os.path.exists(token_file):
        try:
            with open(token_file) as f:
                user_info = json.load(f)
            st.session_state.user = user_info
            st.session_state.current_token = token  # simpan token di session
            saved = load_user(user_info["email"])
            if saved:
                st.session_state.theme = saved.get("theme", "dark")
                if saved.get("sessions"):
                    st.session_state.sessions = saved["sessions"]
                    st.session_state.active_id = saved.get("active_id")
            st.session_state.data_loaded = True
            restore_images_from_messages()
            # JANGAN clear query params — biarkan token tetap di URL
            st.rerun()
        except: pass

# Load data setelah login
if st.session_state.user and not st.session_state.data_loaded:
    saved = load_user(st.session_state.user["email"])
    if saved:
        st.session_state.theme = saved.get("theme", "dark")
        if saved.get("sessions") and not st.session_state.sessions:
            st.session_state.sessions = saved["sessions"]
            st.session_state.active_id = saved.get("active_id")
    st.session_state.data_loaded = True
    restore_images_from_messages()

# Refresh colors after theme load
C = get_colors(st.session_state.theme)

# ─────────────────────────────────────────────
# GLOBAL CSS
# ─────────────────────────────────────────────
st.markdown(f"""
<style>
* {{ font-family: ui-sans-serif,-apple-system,system-ui,"Segoe UI",sans-serif !important; box-sizing: border-box; }}

/* Global BG */
.stApp, [data-testid="stAppViewContainer"],
[data-testid="stAppViewContainer"] > section,
section[data-testid="stMain"],
[data-testid="stMainBlockContainer"],
[data-testid="stBottom"],
[data-testid="stBottom"] > div {{
    background: {C['bg']} !important;
}}

/* Sidebar */
section[data-testid="stSidebar"],
section[data-testid="stSidebar"] > div,
section[data-testid="stSidebar"] > div > div,
section[data-testid="stSidebar"] > div > div > div,
[data-testid="stSidebarContent"],
[data-testid="stSidebarUserContent"],
[data-testid="stSidebarUserContent"] > div,
[data-testid="stSidebarUserContent"] > div > div {{
    background: {C['sidebar_bg']} !important;
    box-shadow: none !important;
}}
section[data-testid="stSidebar"] {{
    border-right: 1px solid {C['border']} !important;
}}

/* Zero padding sidebar */
section[data-testid="stSidebar"] > div,
section[data-testid="stSidebar"] > div > div,
[data-testid="stSidebarContent"],
[data-testid="stSidebarUserContent"],
[data-testid="stSidebarUserContent"] > div {{
    padding-top: 0 !important;
    margin-top: 0 !important;
}}

/* Hide Streamlit collapse button */
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapseButton"] {{
    display: none !important;
}}

/* Sidebar buttons */
section[data-testid="stSidebar"] .stButton > button {{
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    color: {C['text']} !important;
    font-size: 0.875rem !important;
    padding: 7px 12px !important;
    border-radius: 8px !important;
    width: 100% !important;
    display: flex !important;
    align-items: center !important;
    justify-content: flex-start !important;
    text-align: left !important;
    min-height: 36px !important;
}}
section[data-testid="stSidebar"] .stButton > button:hover {{
    background: {C['hover']} !important;
}}
section[data-testid="stSidebar"] .stButton > button p,
section[data-testid="stSidebar"] .stButton > button span {{
    margin: 0 !important;
    text-align: left !important;
    color: inherit !important;
    width: 100% !important;
}}

/* Chat messages */
[data-testid="stChatMessage"] {{
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
}}
[data-testid="stChatMessageAvatarUser"],
[data-testid="stChatMessageAvatarAssistant"] {{
    display: none !important;
}}
[data-testid="stChatMessage"] [data-testid="stMarkdownContainer"] {{
    font-size: 0.9rem !important;
    line-height: 1.75 !important;
    color: {C['text']} !important;
    background: transparent !important;
}}

/* Main content */
[data-testid="stMainBlockContainer"] {{
    max-width: 800px !important;
    margin: 0 auto !important;
    padding: 0 16px 120px !important;
    overflow-y: visible !important;
}}
[data-testid="stMainBlockContainer"] p,
[data-testid="stMainBlockContainer"] li,
[data-testid="stMainBlockContainer"] h1,
[data-testid="stMainBlockContainer"] h2,
[data-testid="stMainBlockContainer"] h3 {{
    color: {C['text']} !important;
}}

/* Chat input */
div[data-testid="stChatInputContainer"] {{
    border: 1px solid {C['border']} !important;
    background: {C['input_bg']} !important;
    border-radius: 16px !important;
}}
[data-testid="stChatInput"] textarea {{
    background: {C['input_bg']} !important;
    color: {C['text']} !important;
    font-size: 0.9rem !important;
}}
[data-testid="stChatInput"] textarea::placeholder {{
    color: {C['text_muted']} !important;
}}
[data-testid="stChatInputContainer"] textarea:focus {{
    box-shadow: none !important;
    outline: none !important;
}}

footer, #MainMenu {{ visibility: hidden !important; }}

/* Divider */
hr {{ border-color: {C['border']} !important; }}

/* ── BASE FONT — konsisten semua ukuran ── */
[data-testid="stMarkdownContainer"] *,
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stMarkdownContainer"] span,
[data-testid="stMarkdownContainer"] div {{
    font-size: 0.95rem !important;
    line-height: 1.8 !important;
}}
[data-testid="stMarkdownContainer"] strong,
[data-testid="stMarkdownContainer"] b {{
    font-size: inherit !important;
}}
[data-testid="stMarkdownContainer"] h1 {{ font-size: 1.3rem !important; }}
[data-testid="stMarkdownContainer"] h2 {{ font-size: 1.15rem !important; }}
[data-testid="stMarkdownContainer"] h3 {{ font-size: 1.05rem !important; }}

/* ── MOBILE ── */
@media (max-width: 768px) {{

    [data-testid="stMainBlockContainer"] {{
        max-width: 100% !important;
        padding: 12px 16px 120px !important;
    }}

    /* Font konsisten di mobile — semua elemen sama */
    [data-testid="stMarkdownContainer"],
    [data-testid="stMarkdownContainer"] *,
    [data-testid="stMarkdownContainer"] p,
    [data-testid="stMarkdownContainer"] li,
    [data-testid="stMarkdownContainer"] span,
    [data-testid="stMarkdownContainer"] strong,
    [data-testid="stMarkdownContainer"] b,
    [data-testid="stMarkdownContainer"] em {{
        font-size: 1rem !important;
        line-height: 1.85 !important;
    }}
    [data-testid="stMarkdownContainer"] h1 {{ font-size: 1.25rem !important; }}
    [data-testid="stMarkdownContainer"] h2 {{ font-size: 1.1rem !important; }}
    [data-testid="stMarkdownContainer"] h3 {{ font-size: 1rem !important; font-weight: 700 !important; }}

    /* List lebih rapi */
    [data-testid="stMarkdownContainer"] ul,
    [data-testid="stMarkdownContainer"] ol {{
        padding-left: 20px !important;
        margin: 6px 0 !important;
    }}
    [data-testid="stMarkdownContainer"] li {{
        margin-bottom: 4px !important;
    }}

    /* Code block */
    [data-testid="stMarkdownContainer"] code {{
        font-size: 0.85rem !important;
        padding: 2px 6px !important;
        border-radius: 4px !important;
        background: rgba(255,255,255,0.08) !important;
    }}
    [data-testid="stMarkdownContainer"] pre {{
        font-size: 0.82rem !important;
        overflow-x: auto !important;
        padding: 12px !important;
        border-radius: 8px !important;
    }}

    /* Chat input */
    div[data-testid="stChatInputContainer"] {{
        border-radius: 26px !important;
        margin: 0 6px 8px !important;
    }}
    [data-testid="stChatInput"] textarea {{
        font-size: 16px !important;
        line-height: 1.5 !important;
    }}

    /* Pesan lebih bernapas */
    [data-testid="stChatMessage"] {{
        padding: 10px 0 !important;
    }}

    /* Bubble user */
    .navy-pill {{
        max-width: 82% !important;
        font-size: 1rem !important;
        line-height: 1.7 !important;
        padding: 12px 16px !important;
    }}
}}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# LOGIN PAGE
# ─────────────────────────────────────────────
def show_login():
    st.markdown(f"""
    <style>
    [data-testid="stSidebar"] {{ display: none !important; }}

    /* Background full screen kipmd.png (desktop) / kipmm.png (mobile) */
    [data-testid="stAppViewContainer"],
    section[data-testid="stMain"] {{
        background: url('https://raw.githubusercontent.com/kipmuniversitaspancasila-commits/KIPMSIGMA/main/kipmd.png') center/cover no-repeat fixed !important;
        min-height: 100vh !important;
    }}

    /* Hapus pseudo-element kiri */
    section[data-testid="stMain"]::before {{
        display: none !important;
    }}

    /* Form container — transparan glass, compact, geser 1cm dari kanan */
    [data-testid="stMainBlockContainer"] {{
        max-width: 300px !important;
        margin: 1.5vh 74px 0 auto !important;
        padding: 8px 18px 16px !important;
        position: relative;
        z-index: 1;
        min-height: unset !important;
        height: fit-content !important;
        background: rgba(5, 8, 20, 0.60) !important;
        backdrop-filter: blur(20px) saturate(1.4) !important;
        -webkit-backdrop-filter: blur(20px) saturate(1.4) !important;
        border: 1px solid rgba(255,255,255,0.10) !important;
        border-radius: 20px !important;
        box-shadow: 0 8px 40px rgba(0,0,0,0.5) !important;
    }}
    @media(max-width: 768px) {{
        [data-testid="stMainBlockContainer"] {{
            margin: 5vh auto 0 auto !important;
            max-width: 88% !important;
            padding: 20px 20px 28px !important;
            backdrop-filter: blur(20px) !important;
            border-radius: 20px !important;
            border: 1px solid rgba(255,255,255,0.12) !important;
            box-shadow: 0 8px 40px rgba(0,0,0,0.5) !important;
        }}
    }}

    /* Sembunyikan header toolbar Streamlit bawaan */
    header[data-testid="stHeader"] {{
        display: none !important;
    }}
    #MainMenu {{
        display: none !important;
    }}

    /* Mobile — pastikan background terlihat di atas & bawah kotak */
    @media(max-width: 768px) {{
        [data-testid="stAppViewContainer"],
        section[data-testid="stMain"] {{
            background: url('https://raw.githubusercontent.com/kipmuniversitaspancasila-commits/KIPMSIGMA/main/kipmm.png') center top/cover no-repeat fixed !important;
        }}
        /* Kotak login beri ruang logo KIPM di atas */
        [data-testid="stMainBlockContainer"] {{
            margin-top: 75px !important;
        }}
    }}

    /* Glass card */
    .stTabs, [data-testid="stVerticalBlock"] {{
        background: transparent !important;
    }}

    /* Input fields glass style */
    [data-testid="stTextInput"] input {{
        background: rgba(255,255,255,0.06) !important;
        border: 1px solid rgba(255,255,255,0.12) !important;
        border-radius: 12px !important;
        color: #fff !important;
        padding: 12px 16px !important;
        font-size: 0.95rem !important;
        backdrop-filter: blur(10px) !important;
        transition: border 0.2s !important;
    }}
    [data-testid="stTextInput"] input:focus {{
        border: 1px solid {C['gold']} !important;
        box-shadow: 0 0 0 2px rgba(245,194,66,0.15) !important;
        outline: none !important;
    }}
    [data-testid="stTextInput"] input::placeholder {{ color: rgba(255,255,255,0.35) !important; }}
    [data-testid="stTextInput"] label {{ color: rgba(255,255,255,0.6) !important; font-size: 0.82rem !important; }}

    /* Masuk button */
    [data-testid="stMainBlockContainer"] .stButton > button {{
        background: linear-gradient(135deg, {C['gold']}, #e0a820) !important;
        color: #000 !important;
        font-weight: 700 !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 12px !important;
        font-size: 0.95rem !important;
        letter-spacing: 0.5px !important;
        transition: opacity 0.2s, transform 0.1s !important;
        box-shadow: 0 4px 20px rgba(245,194,66,0.3) !important;
    }}
    [data-testid="stMainBlockContainer"] .stButton > button:hover {{
        opacity: 0.92 !important; transform: translateY(-1px) !important;
    }}

    /* Tabs glass */
    [data-testid="stTabs"] [role="tablist"] {{
        background: rgba(255,255,255,0.05) !important;
        border-radius: 12px !important;
        padding: 4px !important;
        border: 1px solid rgba(255,255,255,0.08) !important;
        gap: 2px !important;
    }}
    [data-testid="stTabs"] button[role="tab"] {{
        border-radius: 9px !important;
        color: rgba(255,255,255,0.5) !important;
        font-size: 0.85rem !important;
        padding: 7px 12px !important;
        border: none !important;
        background: transparent !important;
    }}
    [data-testid="stTabs"] button[role="tab"][aria-selected="true"] {{
        background: rgba(245,194,66,0.15) !important;
        color: {C['gold']} !important;
        font-weight: 600 !important;
    }}
    [data-testid="stTabs"] [role="tabpanel"] {{
        background: rgba(255,255,255,0.03) !important;
        border-radius: 16px !important;
        border: 1px solid rgba(255,255,255,0.08) !important;
        padding: 20px 16px !important;
        margin-top: 8px !important;
        backdrop-filter: blur(10px) !important;
    }}

    /* Alert colors */
    [data-testid="stAlert"] {{ border-radius: 10px !important; }}

    /* Animasi background */
    @keyframes gradMove {{
        0% {{ background-position: 0% 50%; }}
        50% {{ background-position: 100% 50%; }}
        100% {{ background-position: 0% 50%; }}
    }}
    </style>
    """, unsafe_allow_html=True)

    # Background sudah di-set via CSS (URL GitHub raw)

    # Inject logo KIPM khusus mobile — hanya saat halaman login
    components.html(f"""
<script>
(function() {{
    var pd = window.parent.document;

    // Sembunyikan Fork bar & toolbar GitHub Streamlit
    var forkStyle = pd.getElementById('hide-fork-bar');
    if (!forkStyle) {{
        var fs = pd.createElement('style');
        fs.id = 'hide-fork-bar';
        fs.textContent = `
            /* Fork bar & GitHub badge */
            .viewerBadge_container__r5tak,
            .viewerBadge_link__qRIco,
            [class*="viewerBadge"],
            [class*="styles_viewerBadge"],
            #MainMenu,
            footer,
            /* Streamlit toolbar atas */
            [data-testid="stToolbar"],
            [data-testid="stDecoration"],
            [data-testid="stStatusWidget"],
            header[data-testid="stHeader"],
            /* Tombol deploy/fork */
            .stDeployButton,
            [kind="header"],
            div[data-testid="collapsedControl"] {{
                display: none !important;
                visibility: hidden !important;
                height: 0 !important;
                overflow: hidden !important;
            }}
        `;
        pd.head.appendChild(fs);
    }}

    if (pd.getElementById('kipm-mobile-logo')) return;

    var s = pd.createElement('style');
    s.id = 'kipm-mobile-logo-style';
    s.textContent = `
        #kipm-mobile-logo {{
            display: none;
            text-align: center;
            padding: 14px 0 10px;
            position: fixed;
            top: 0; left: 0; right: 0;
            z-index: 10;
            pointer-events: none;
        }}
        #kipm-mobile-logo img {{
            width: 80px; height: 80px;
            object-fit: contain;
            filter: drop-shadow(0 2px 12px rgba(0,0,0,0.6));
        }}
        #kipm-mobile-logo .kipm-name {{
            font-size: 0.7rem;
            color: rgba(255,255,255,0.7);
            letter-spacing: 2px;
            font-family: sans-serif;
            margin-top: 4px;
        }}
        @media(max-width: 768px) {{
            #kipm-mobile-logo {{ display: block !important; }}
        }}
    `;
    pd.head.appendChild(s);

    var div = pd.createElement('div');
    div.id = 'kipm-mobile-logo';
    div.innerHTML = `
        <img src="https://raw.githubusercontent.com/kipmuniversitaspancasila-commits/KIPMSIGMA/main/Mate%20KIPM%20LOGO.png"
             onerror="this.style.display='none'"
             style="width:80px;height:80px;object-fit:contain;">
        <div class="kipm-name">KIPM-UP</div>
    `;
    pd.body.appendChild(div);
}})();
</script>
""", height=0)
    st.markdown('''
        <div style="text-align:center;margin:0 0 10px;">
            <div style="font-size:2.8rem;font-weight:900;letter-spacing:5px;color:#ffffff;font-family:sans-serif;line-height:1.2;">
                SIGMA <span style="color:#F5C242;">Σ</span>
            </div>
            <div class="sigma-tagline" style="font-size:0.65rem;color:rgba(255,255,255,0.5);letter-spacing:2px;margin-top:4px;font-family:sans-serif;">
                Strategic Intelligence & Global Market Analysis
            </div>
        </div>
        <style>
            @media(min-width: 769px) { .sigma-tagline { display: none !important; } }
        </style>
    ''', unsafe_allow_html=True)
    tab1, tab2, tab3 = st.tabs(["🔑 Sign In", "📝 Sign Up", "🌐 Google"])

    with tab1:
        uname = st.text_input("Username", key="li_user", placeholder="Masukkan username")
        pwd   = st.text_input("Password", key="li_pwd",  type="password", placeholder="Masukkan password")
        if st.button("Masuk", key="btn_login", use_container_width=True):
            if uname and pwd:
                info = login_user(uname.strip(), pwd)
                if info:
                    token = str(uuid.uuid4()).replace("-","")
                    with open(os.path.join(DATA_DIR, f"token_{token}.json"), "w") as f:
                        json.dump(info, f)
                    st.query_params["sigma_token"] = token
                    st.session_state.user = info
                    st.session_state.current_token = token
                    st.session_state.data_loaded = False
                    st.rerun()
                else:
                    st.error("Username atau password salah")
            else:
                st.warning("Isi username dan password")

    with tab2:
        rname  = st.text_input("Nama Tampil", key="rg_name", placeholder="Nama lengkap kamu")
        runame = st.text_input("Username", key="rg_user", placeholder="username (huruf/angka)")
        rpwd   = st.text_input("Password", key="rg_pwd",  type="password", placeholder="min. 6 karakter")
        rpwd2  = st.text_input("Ulangi Password", key="rg_pwd2", type="password", placeholder="ulangi password")
        if st.button("Daftar Sekarang", key="btn_register", use_container_width=True):
            if not all([rname, runame, rpwd, rpwd2]):
                st.warning("Lengkapi semua field")
            elif rpwd != rpwd2:
                st.error("Password tidak cocok")
            elif len(rpwd) < 6:
                st.error("Password minimal 6 karakter")
            else:
                ok, msg = register_user(runame.strip(), rpwd, rname.strip())
                if ok:
                    st.success(f"✅ {msg} — silakan masuk")
                else:
                    st.error(msg)

    with tab3:
        try:
            auth_url = google_auth_url()
            st.markdown(f"""
            <div style="margin-top:8px;">
                <a href="{auth_url}" style="
                    display:flex;align-items:center;justify-content:center;gap:10px;
                    background:rgba(255,255,255,0.95);color:#1a1a1a;border-radius:12px;padding:13px;
                    text-decoration:none;font-size:0.9rem;font-weight:600;
                    border:none;box-shadow:0 4px 15px rgba(0,0,0,0.3);">
                    <svg width="18" height="18" viewBox="0 0 24 24">
                        <path fill="#4285F4" d="M22.56 12.25c0-.78-.07-1.53-.2-2.25H12v4.26h5.92c-.26 1.37-1.04 2.53-2.21 3.31v2.77h3.57c2.08-1.92 3.28-4.74 3.28-8.09z"/>
                        <path fill="#34A853" d="M12 23c2.97 0 5.46-.98 7.28-2.66l-3.57-2.77c-.98.66-2.23 1.06-3.71 1.06-2.86 0-5.29-1.93-6.16-4.53H2.18v2.84C3.99 20.53 7.7 23 12 23z"/>
                        <path fill="#FBBC05" d="M5.84 14.09c-.22-.66-.35-1.36-.35-2.09s.13-1.43.35-2.09V7.07H2.18C1.43 8.55 1 10.22 1 12s.43 3.45 1.18 4.93l3.66-2.84z"/>
                        <path fill="#EA4335" d="M12 5.38c1.62 0 3.06.56 4.21 1.64l3.15-3.15C17.45 2.09 14.97 1 12 1 7.7 1 3.99 3.47 2.18 7.07l3.66 2.84c.87-2.6 3.3-4.53 6.16-4.53z"/>
                    </svg>
                    Lanjutkan dengan Google
                </a>
            </div>
            """, unsafe_allow_html=True)
        except:
            st.info("Google login belum dikonfigurasi di Secrets")

    st.markdown(f"""
    <p style="text-align:center;color:rgba(255,255,255,0.25);font-size:0.72rem;margin-top:24px;line-height:1.6;">
        Dengan masuk, kamu menyetujui penggunaan platform untuk analisa.<br>
        Analisa bersifat <em>do your own research</em> dan disclaimer berlaku. by MarketnMocha
    </p>
    """, unsafe_allow_html=True)
    st.stop()

# ─────────────────────────────────────────────
# GUARD: LOGIN REQUIRED
# ─────────────────────────────────────────────
if st.session_state.user is None:
    show_login()

# Init chat after login
init_chat()
user = st.session_state.user
C = get_colors(st.session_state.theme)


# ─────────────────────────────────────────────
# SEMBUNYIKAN SIDEBAR
# ─────────────────────────────────────────────
st.markdown("""
<style>
section[data-testid="stSidebar"],
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapseButton"] { display: none !important; }

/* Sembunyikan Fork bar, GitHub badge, Streamlit logo di semua halaman */
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"],
.viewerBadge_container__r5tak,
[class*="viewerBadge"],
.stDeployButton,
#MainMenu,
footer,
/* Bar hitam Fork */
[data-testid="stHeader"],
iframe[title="streamlit_analytics"],
div[class*="Toolbar"],
div[class*="toolbar"],
div[class*="ActionButton"],
div[class*="HeaderActionButton"] { display: none !important; }
</style>
""", unsafe_allow_html=True)

_hist_items = ""
for _sesi in st.session_state.sessions:
    _sid = _sesi["id"]
    _is_act = _sid == st.session_state.active_id
    _td = _sesi["title"][:35].replace("'","").replace("`","").replace("\\","").replace('"',"")
    _fw = "700" if _is_act else "400"
    _bg = C['hover'] if _is_act else "transparent"
    # Pakai <a href> langsung — tidak perlu JS event listener
    _hist_items += f"""
(function(){{
    var a=pd.createElement('a');
    a.textContent='{_td}';
    var u=new URL(window.parent.location.href);
    u.searchParams.set('do','sel_{_sid}');
    a.href=u.toString();
    a.style.cssText='display:block;width:100%;padding:12px 18px;font-size:1rem;color:{C["text"]};background:{_bg};font-weight:{_fw};border:none;text-align:left;cursor:pointer;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;text-decoration:none;';
    a.onmouseenter=function(){{this.style.background='{C["hover"]}'}};
    a.onmouseleave=function(){{this.style.background='{_bg}'}};
    h.appendChild(a);
}})();
"""

# Inject menu ke parent document via components.html
components.html(f"""
<script>
(function(){{
var pd=window.parent.document;

// Sembunyikan logo KIPM login saat sudah masuk ke halaman utama
var kipmLogo = pd.getElementById('kipm-mobile-logo');
if (kipmLogo) kipmLogo.style.display = 'none !important';
var kipmStyle = pd.getElementById('kipm-mobile-logo-style');
if (kipmStyle) kipmStyle.remove();

// Hapus elemen lama kalau ada, inject ulang yang baru
['spbtn','spmenu','sphist','spui','sigma-mobile-css'].forEach(function(id){{
    var el=pd.getElementById(id);
    if(el) el.remove();
}});

// CSS
var s=pd.createElement('style');
s.id='sigma-mobile-css';
s.textContent=`
#spbtn{{position:fixed;bottom:16px;left:14px;width:38px;height:38px;border-radius:50%;
    background:{C["sidebar_bg"]};color:{C["text"]};border:1px solid {C["border"]};
    cursor:pointer;font-size:24px;font-weight:300;z-index:99999;
    display:flex;align-items:center;justify-content:center;
    box-shadow:0 2px 10px rgba(0,0,0,0.5);padding:0;line-height:1;}}
#spbtn:hover{{transform:scale(1.1)}}
#spmenu,#sphist{{position:fixed;left:12px;bottom:62px;
    background:{C["sidebar_bg"]};border:1px solid {C["border"]};
    border-radius:16px;box-shadow:0 -4px 24px rgba(0,0,0,0.5);
    z-index:99998;display:none;overflow:hidden;min-width:250px;}}
#sphist{{max-height:55vh;overflow-y:auto;}}
.smi{{display:flex;align-items:center;gap:14px;padding:13px 18px;
    font-size:1rem;color:{C["text"]};cursor:pointer;border:none;
    background:transparent;width:100%;text-align:left;}}
.smi:hover{{background:{C["hover"]}}}
.smico{{width:32px;height:32px;border-radius:8px;display:flex;
    align-items:center;justify-content:center;font-size:16px;
    background:{C["hover"]};flex-shrink:0;}}
.smsp{{border:none;border-top:1px solid {C["border"]};margin:4px 0;}}
.smhd{{padding:8px 18px 4px;font-size:0.68rem;color:{C["text_muted"]};
    font-weight:600;letter-spacing:1px;}}
.smred{{color:#f55!important}}
`;
pd.head.appendChild(s);

// + button
var btn=pd.createElement('button');
btn.id='spbtn';btn.textContent='+';pd.body.appendChild(btn);

// Menu
var m=pd.createElement('div');m.id='spmenu';
m.innerHTML=`
<a class="smi" id="smi-new"><span class="smico">✎</span>Obrolan baru</a>
<button class="smi" id="smi-hist"><span class="smico">☰</span>Riwayat obrolan</button>
<div class="smsp"></div>
<div class="smhd">PENAMPILAN</div>
<a class="smi" id="smi-dark"><span class="smico">🌙</span>Mode Gelap {'✓' if st.session_state.theme=='dark' else ''}</a>
<a class="smi" id="smi-light"><span class="smico">☀️</span>Mode Terang {'✓' if st.session_state.theme=='light' else ''}</a>
<div class="smsp"></div>
<a class="smi smred" id="smi-out"><span class="smico">🚪</span>Keluar</a>
`;
pd.body.appendChild(m);

// Nav function
function nav(params){{
    var u=new URL(window.parent.location.href);
    Object.keys(params).forEach(function(k){{u.searchParams.set(k,params[k])}});
    window.parent.location.href=u.toString();
}}

// History drawer
var h=pd.createElement('div');h.id='sphist';
h.innerHTML='<div class="smhd">RIWAYAT OBROLAN</div>';
{_hist_items}
pd.body.appendChild(h);

btn.onclick=function(e){{e.stopPropagation();m.style.display=m.style.display==='block'?'none':'block';h.style.display='none';}};

// Set href langsung pada menu items
(function(){{
    var u;
    u=new URL(window.parent.location.href); u.searchParams.set('do','newchat');
    pd.getElementById('smi-new').href=u.toString();
    pd.getElementById('smi-new').style.textDecoration='none';
    
    pd.getElementById('smi-hist').onclick=function(){{m.style.display='none';h.style.display=h.style.display==='block'?'none':'block';}};
    
    u=new URL(window.parent.location.href); u.searchParams.set('do','theme_dark');
    pd.getElementById('smi-dark').href=u.toString();
    pd.getElementById('smi-dark').style.textDecoration='none';
    
    u=new URL(window.parent.location.href); u.searchParams.set('do','theme_light');
    pd.getElementById('smi-light').href=u.toString();
    pd.getElementById('smi-light').style.textDecoration='none';
    
    u=new URL(window.parent.location.href); u.searchParams.delete('sigma_token'); u.searchParams.set('do','logout');
    pd.getElementById('smi-out').href=u.toString();
    pd.getElementById('smi-out').style.textDecoration='none';
}})();

pd.addEventListener('click',function(e){{
    if(!btn.contains(e.target)&&!m.contains(e.target))m.style.display='none';
    if(!btn.contains(e.target)&&!h.contains(e.target)&&!m.contains(e.target))h.style.display='none';
}});
}})();
</script>
""", height=0)

# Handle actions
if "do" in st.query_params:
    _do = st.query_params.get("do", "")
    _tok = st.query_params.get("sigma_token", st.session_state.get("current_token", ""))
    if _do == "logout":
        if _tok:
            try: os.remove(os.path.join(DATA_DIR, f"token_{_tok}.json"))
            except: pass
        st.session_state.clear(); st.query_params.clear(); st.rerun()
    elif _do == "theme_dark":
        st.session_state.theme = "dark"
        st.query_params["do"] = ""; st.rerun()
    elif _do == "theme_light":
        st.session_state.theme = "light"
        st.query_params["do"] = ""; st.rerun()
    elif _do == "newchat":
        ns = new_session()
        st.session_state.sessions.insert(0, ns)
        st.session_state.active_id = ns["id"]
        st.query_params["do"] = ""; st.rerun()
    elif _do.startswith("sel_"):
        _sid = _do[4:]
        st.session_state.active_id = _sid
        st.query_params["do"] = ""; st.rerun()

# ─────────────────────────────────────────────
# MAIN CHAT
# ─────────────────────────────────────────────
active = get_active()

# Build history JS untuk drawer — inject items tanpa st.button
_hist_items = ""
for sesi in st.session_state.sessions:
    sid = sesi["id"]
    is_active = sid == st.session_state.active_id
    title_d = sesi["title"][:35].replace("'", "\\'").replace("`","")
    fw = "600" if is_active else "400"
    bg = C['hover'] if is_active else "transparent"
    _hist_items += f"""
    (function() {{
        var hi = document.createElement('button');
        hi.textContent = '{title_d}';
        hi.dataset.sid = '{sid}';
        hi.style.cssText = 'display:block;width:100%;padding:11px 16px;font-size:0.95rem;color:{C["text"]};background:{bg};font-weight:{fw};border:none;text-align:left;cursor:pointer;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;';
        hi.onmouseenter = function(){{this.style.background='{C["hover"]}'}};
        hi.onmouseleave = function(){{this.style.background='{bg}'}};
        hi.onclick = function(){{
            var url = new URL(window.parent.location.href);
            url.searchParams.set('do', 'sel_{sid}');
            window.parent.location.href = url.toString();
        }};
        drawer.appendChild(hi);
    }})();
"""

# Header
if not active["messages"][1:]:
    uname = user.get("name", "").split()[0] if user.get("name") else "Trader"
    st.markdown(f"""
    <div style="text-align:center;padding:10vh 0 2rem;">
        <h1 style="margin:0;font-size:1.8rem;font-weight:700;color:{C['text']};">Halo, {uname} 👋</h1>
        <p style="margin:8px 0 0;color:{C['text_muted']};font-size:0.9rem;">Ada yang bisa SIGMA bantu analisa hari ini?</p>
    </div>
    """, unsafe_allow_html=True)

# Chat history
for i, msg in enumerate(active["messages"][1:]):
    with st.chat_message(msg["role"]):
        # Gunakan field "display" (prompt bersih) kalau ada, fallback ke "content"
        raw = msg.get("display") or msg["content"]
        display = raw
        if "Pertanyaan:" in display:
            display = display.split("Pertanyaan:")[-1].strip()
        # Bersihkan sisa marker context kalau ada di pesan lama
        if "=== DATA PASAR REAL-TIME ===" in display:
            parts = display.split("===========================")
            display = parts[-1].strip() if len(parts) > 1 else display
        if msg["role"] == "user" and msg.get("img_b64"):
            st.markdown(f'<img src="data:{msg.get("img_mime","image/jpeg")};base64,{msg["img_b64"]}" style="max-width:100%;max-height:240px;border-radius:10px;margin-bottom:6px;display:block;">', unsafe_allow_html=True)
        st.markdown(display)
        # Tampilkan tombol download untuk pesan analisa di history
        if msg["role"] == "assistant" and msg.get("is_analisa"):
            _col1, _col2 = st.columns(2)
            with _col1:
                try:
                    _xb = create_excel_download(msg["content"])
                    if _xb:
                        st.download_button("⬇️ Excel", _xb,
                            file_name=f"sigma_{msg.get('id','analisa')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"dl_xlsx_{i}", use_container_width=True)
                except Exception:
                    pass
            with _col2:
                try:
                    _wb = create_word_download(msg["content"])
                    if _wb:
                        st.download_button("⬇️ Word", _wb,
                            file_name=f"sigma_{msg.get('id','analisa')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_docx_{i}", use_container_width=True)
                except Exception:
                    pass

# Chat input
try:
    result = st.chat_input(
        "Tanya SIGMA... DYOR - bukan financial advice.",
        accept_file="multiple",
        file_type=["pdf", "png", "jpg", "jpeg", "xlsx", "xls", "docx", "doc", "csv"]
    )
except TypeError:
    result = st.chat_input("Tanya SIGMA...")

prompt = None
file_obj = None

if result is not None:
    if hasattr(result, 'text'):
        prompt = (result.text or "").strip()
        files = getattr(result, 'files', None) or []
        if files: file_obj = files[0]
    elif isinstance(result, str):
        prompt = result.strip()

    if file_obj:
        raw = file_obj.read()
        fname = file_obj.name.lower()
        ftype = file_obj.type

        # Excel
        if fname.endswith((".xlsx", ".xls")) or "spreadsheet" in ftype or "excel" in ftype:
            excel_text = read_excel_file(raw, file_obj.name)
            st.session_state.pdf_data = (excel_text, file_obj.name)
            st.session_state.img_data = None
        # Word
        elif fname.endswith((".docx", ".doc")) or "word" in ftype or "document" in ftype:
            word_text = read_word_file(raw, file_obj.name)
            st.session_state.pdf_data = (word_text, file_obj.name)
            st.session_state.img_data = None
        # CSV
        elif fname.endswith(".csv") or "csv" in ftype:
            try:
                csv_text = raw.decode("utf-8", errors="replace")
                st.session_state.pdf_data = (f"[CSV: {file_obj.name}]\n{csv_text[:30000]}", file_obj.name)
            except Exception:
                st.session_state.pdf_data = (f"[CSV: {file_obj.name} — gagal dibaca]", file_obj.name)
            st.session_state.img_data = None
        elif ftype == "application/pdf":
            doc = fitz.open(stream=raw, filetype="pdf")

            # ── Smart PDF extractor — cari halaman keuangan, skip halaman tidak relevan ──
            FINANCE_KEYWORDS = [
                # Laporan utama
                "laba rugi", "neraca", "arus kas", "ekuitas", "balance sheet",
                "income statement", "cash flow", "profit", "revenue", "pendapatan",
                "beban", "laba bersih", "net income", "total aset", "liabilitas",
                # Rasio perbankan
                "nim", "npl", "ldr", "car", "roa", "roe", "bopo", "cir",
                "net interest margin", "non performing", "loan to deposit",
                "capital adequacy", "cost to income",
                # Fundamental umum
                "eps", "earning per share", "laba per saham", "dividen", "dividend",
                "payout", "book value", "nilai buku", "kredit", "pembiayaan",
                "dana pihak ketiga", "dpk", "modal", "retained earning",
                "saldo laba", "gross profit", "operating profit",
                # Tabel angka — deteksi halaman berisi banyak angka
                "rp ", "rp.", "000.000", "miliar", "triliun", "%",
            ]

            relevant_pages = []
            all_pages_text = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                all_pages_text.append((page_num + 1, page_text))
                page_lower = page_text.lower()
                # Hitung berapa keyword keuangan yang ditemukan di halaman ini
                score = sum(1 for kw in FINANCE_KEYWORDS if kw in page_lower)
                if score >= 3:  # minimal 3 keyword keuangan
                    relevant_pages.append((score, page_num + 1, page_text))

            # Urutkan dari yang paling relevan, ambil top 40 halaman
            relevant_pages.sort(key=lambda x: x[0], reverse=True)
            top_pages = relevant_pages[:40]
            # Urutkan ulang berdasarkan nomor halaman agar berurutan
            top_pages.sort(key=lambda x: x[1])

            if top_pages:
                extracted = f"[PDF: {file_obj.name} | {len(doc)} halaman | {len(top_pages)} halaman relevan diekstrak]\n\n"
                for score, pnum, ptext in top_pages:
                    # Bersihkan whitespace berlebih
                    clean = " ".join(ptext.split())
                    extracted += f"--- Halaman {pnum} ---\n{clean[:2000]}\n\n"
                # Batas total ~60.000 karakter agar tidak overflow context window
                final_text = extracted[:60000]
            else:
                # Fallback: tidak ada halaman keuangan terdeteksi, ambil 20 halaman pertama
                fallback = f"[PDF: {file_obj.name} | {len(doc)} halaman | fallback: 20 halaman pertama]\n\n"
                for pnum, ptext in all_pages_text[:20]:
                    clean = " ".join(ptext.split())
                    fallback += f"--- Halaman {pnum} ---\n{clean[:1500]}\n\n"
                final_text = fallback[:40000]

            # Inject data pasar real-time untuk melengkapi data yang tidak ada di PDF
            try:
                fundamental_ctx = build_fundamental_context(final_text)
                if fundamental_ctx:
                    final_text = final_text + fundamental_ctx
            except Exception:
                pass

            st.session_state.pdf_data = (final_text, file_obj.name)
            st.session_state.img_data = None
        else:
            b64 = base64.b64encode(raw).decode()
            mime = "image/png" if file_obj.name.endswith(".png") else "image/jpeg"
            st.session_state.img_data = (b64, mime, file_obj.name)
            st.session_state.pdf_data = None

    if not prompt and (file_obj or st.session_state.img_data or st.session_state.pdf_data):
        if st.session_state.pdf_data:
            prompt = f"Tolong analisa laporan keuangan: {st.session_state.pdf_data[1]}"
        else:
            prompt = "Tolong analisa file yang saya kirim"

if prompt:
    img_data = st.session_state.img_data
    pdf_data = st.session_state.pdf_data
    # Clear SETELAH diambil nilainya
    st.session_state.img_data = None
    st.session_state.pdf_data = None

    # Hitung estimasi token — 1 token ≈ 4 karakter
    # llama-3.3-70b context window = 128k token ≈ 512k karakter
    # System prompt baru ~3k karakter, sisakan ruang untuk output
    MAX_PDF_CHARS = 80000

    full_prompt = prompt
    if img_data:
        full_prompt = f"[Gambar: {img_data[2]}]\n\nPertanyaan: {prompt}"
    elif pdf_data:
        pdf_text = pdf_data[0]
        if len(pdf_text) > MAX_PDF_CHARS:
            pdf_text = pdf_text[:MAX_PDF_CHARS] + "\n\n[...konten terpotong, analisa berdasarkan data di atas]"
        full_prompt = f"{pdf_text}\n\nPertanyaan: {prompt}"

    # ── Deteksi intent: inject data sesuai jenis permintaan ──
    if not img_data:
        _p = prompt.lower()

        # Deteksi perintah analisa fundamental tanpa PDF
        _fundamental_keywords = ["fundamental", "laporan keuangan", "keuangan", "valuasi",
                                  "ipo", "historis", "proyeksi", "per ", "pbv", "roe", "roa"]
        _teknikal_keywords = ["analisa", "analisis", "entry", "sl", "tp", "target",
                               "stop loss", "beli", "jual", "hold", "chart", "teknikal",
                               "bias", "setup", "volume", "bandar", "bandarmologi",
                               "support", "resistance", "breakout", "breakdown"]
        # Deteksi ticker — cek huruf besar DAN kecil
        _prompt_upper = prompt.upper()
        # Keyword umum — semua prompt ini tetap dikirim ke Groq
        _general_keywords = ["buatkan", "buat", "word", "excel", "download", "file",
                              "ringkasan", "summary", "jelaskan", "ceritakan", "apa",
                              "bagaimana", "kenapa", "kapan", "berapa", "tolong",
                              "coba", "bisa", "mau", "ingin", "minta", "hai", "halo",
                              "selamat", "terima", "makasih", "thanks"]
        _has_ticker = bool(re.search(r'\b[A-Z]{4}\b', _prompt_upper))
        _is_fundamental_cmd = _has_ticker and any(k in _p for k in _fundamental_keywords)
        _is_teknikal = _has_ticker or any(k in _p for k in _teknikal_keywords)
        _is_general = not _is_fundamental_cmd and not _is_teknikal and any(k in _p for k in _general_keywords)

        if _is_fundamental_cmd and not pdf_data:
            # Perintah analisa fundamental tanpa PDF — tarik data lengkap dari yfinance
            tickers_found = re.findall(r'\b([A-Z]{4})\b', _prompt_upper)
            if tickers_found:
                _skip_words = {"YANG","ATAU","DARI","PADA","UNTUK","DENGAN","SAHAM",
                               "SAYA","MINTA","TOLONG","ANALISA","ANALISIS","MOHON",
                               "BBRI","BBCA","BMRI","TLKM"}  # jangan skip ticker valid
                tickers_found = [t for t in tickers_found if t not in _skip_words or t in
                                 {"BBRI","BBCA","BMRI","TLKM","ASII","GOTO","BRIS","UNVR",
                                  "ANTM","PTBA","ADRO","EXCL","SMGR","KLBF","SIDO","CPIN"}]
            if tickers_found:
                _ticker = tickers_found[0]
                # Fetch dengan timeout agar tidak block Groq call
                fund_ctx = ""
                try:
                    import threading
                    _result = [None]
                    def _fetch():
                        _result[0] = fetch_full_fundamental(_ticker)
                    t = threading.Thread(target=_fetch)
                    t.start()
                    t.join(timeout=15)  # max 15 detik untuk fetch data
                    fund_ctx = _result[0] or f"[Data fetch timeout — analisa dari knowledge model]"
                except Exception as _fe:
                    fund_ctx = f"[Gagal fetch: {_fe} — analisa dari knowledge model]"
                full_prompt = (
                    f"{fund_ctx}\n\n"
                    f"Perintah: {prompt}\n\n"
                    f"Instruksi: Buat analisa fundamental lengkap FORMAT ANALISA FUNDAMENTAL "
                    f"untuk saham {_ticker}. Tren 3 tahun terakhir dan proyeksi 1-2 tahun ke depan. "
                    f"Jika data di atas kosong/gagal, TETAP buat analisa lengkap dari knowledge kamu."
                )
        elif _is_teknikal:
            # Perintah teknikal — inject market context biasa
            try:
                mkt_ctx = get_market_context(_prompt_upper)
                if mkt_ctx:
                    full_prompt = (
                        f"=== DATA PASAR REAL-TIME ===\n{mkt_ctx}\n"
                        f"===========================\n\n{full_prompt}"
                    )
            except Exception:
                pass
        # else: perintah umum/general — langsung kirim ke Groq tanpa inject data
        # full_prompt tetap = prompt asli, Groq tetap menjawab

    if active["title"] == "Obrolan Baru":
        active["title"] = prompt[:40] + ("..." if len(prompt) > 40 else "")

    # Simpan gambar di dalam message agar tetap ada setelah refresh
    # PENTING: "content" diisi prompt bersih (tanpa market context) agar bubble user tetap rapi
    # full_prompt (dengan context) hanya dipakai saat API call ke Groq, tidak disimpan ke history
    user_msg = {"role": "user", "content": full_prompt, "display": prompt}
    if img_data:
        user_msg["img_b64"] = img_data[0]
        user_msg["img_mime"] = img_data[1]
        # Juga simpan di session_state untuk render langsung
        thumb_idx = len(active["messages"]) - 1
        st.session_state[f"thumb_{active['id']}_{thumb_idx}"] = (img_data[0], img_data[1])

    active["messages"].append(user_msg)
    with st.chat_message("user"):
        if img_data:
            st.markdown(f'<img src="data:{img_data[1]};base64,{img_data[0]}" style="max-width:100%;max-height:240px;border-radius:10px;margin-bottom:6px;display:block;">', unsafe_allow_html=True)
        elif pdf_data:
            st.markdown(f'''<div style="display:inline-flex;align-items:center;gap:10px;background:rgba(255,255,255,0.08);border:1px solid rgba(255,255,255,0.15);border-radius:10px;padding:10px 14px;margin-bottom:8px;">
                <span style="font-size:1.4rem;">📄</span>
                <div>
                    <div style="font-size:0.85rem;font-weight:600;color:#ececec;">{pdf_data[1]}</div>
                    <div style="font-size:0.75rem;color:#8e8ea0;">PDF · Laporan Keuangan</div>
                </div>
            </div>''', unsafe_allow_html=True)
        st.markdown(prompt)

    try:
        groq_client = Groq(api_key=st.secrets["GROQ_API_KEY"])
        with st.chat_message("assistant"):
            with st.spinner("SIGMA menganalisis..."):
                if img_data:
                    res = groq_client.chat.completions.create(
                        model="meta-llama/llama-4-scout-17b-16e-instruct",
                        messages=[
                            {"role": "system", "content": "Kamu SIGMA, analis chart. Analisa gambar langsung. Jawab Bahasa Indonesia."},
                            {"role": "user", "content": [
                                {"type": "image_url", "image_url": {"url": f"data:{img_data[1]};base64,{img_data[0]}"}},
                                {"type": "text", "text": prompt}
                            ]}
                        ],
                        max_tokens=2048
                    )
                else:
                    # Bersihkan field "display" — Groq hanya terima "role" dan "content"
                    groq_messages = [
                        {"role": m["role"], "content": m["content"]}
                        for m in active["messages"]
                    ]
                    # Deteksi apakah pesan terakhir adalah PDF — pakai model context lebih besar
                    last_content = groq_messages[-1]["content"] if groq_messages else ""
                    is_pdf_msg = "[PDF:" in last_content and "Halaman" in last_content
                    if is_pdf_msg:
                        # Untuk PDF: kirim hanya system prompt + pesan PDF saja
                        groq_messages_send = [
                            groq_messages[0],   # system prompt
                            groq_messages[-1],  # pesan PDF user
                        ]
                        pdf_content = groq_messages_send[-1]["content"]
                        if len(pdf_content) > 80000:
                            groq_messages_send[-1] = {
                                "role": "user",
                                "content": pdf_content[:80000] + "\n\n[Data terpotong — analisa berdasarkan data di atas]"
                            }
                        model_to_use = "llama-3.3-70b-versatile"
                        max_tok = 4096
                        st.caption(f"📄 Memproses PDF — {len(groq_messages_send[-1]['content']):,} karakter dikirim ke model...")
                    else:
                        groq_messages_send = groq_messages
                        model_to_use = "llama-3.3-70b-versatile"
                        max_tok = 2048
                    res = groq_client.chat.completions.create(
                        model=model_to_use,
                        messages=groq_messages_send,
                        temperature=0.7,
                        max_tokens=max_tok
                    )
                ans = res.choices[0].message.content
            st.markdown(ans)

            # ── Tombol download jika output adalah analisa ──
            _is_analisa = any(k in ans for k in ["TRADE PLAN", "ANALISA FUNDAMENTAL", "PROFITABILITAS", "VERDICT"])
            if _is_analisa:
                col1, col2 = st.columns(2)
                with col1:
                    try:
                        xlsx_bytes = create_excel_download(ans)
                        if xlsx_bytes:
                            st.download_button(
                                label="⬇️ Download Excel",
                                data=xlsx_bytes,
                                file_name=f"sigma_analisa_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True
                            )
                    except Exception:
                        pass
                with col2:
                    try:
                        docx_bytes = create_word_download(ans)
                        if docx_bytes:
                            st.download_button(
                                label="⬇️ Download Word",
                                data=docx_bytes,
                                file_name=f"sigma_analisa_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    except Exception:
                        pass

        active["messages"].append({
            "role": "assistant",
            "content": ans,
            "is_analisa": any(k in ans for k in ["TRADE PLAN", "ANALISA FUNDAMENTAL", "PROFITABILITAS", "VERDICT"])
        })
    except Exception as e:
        err_msg = str(e)
        if "rate_limit" in err_msg.lower():
            st.error("⚠️ Rate limit Groq tercapai — coba lagi dalam beberapa detik.")
        elif "context" in err_msg.lower() or "token" in err_msg.lower():
            st.error("⚠️ Dokumen terlalu besar untuk diproses sekaligus. Coba kirim bagian tertentu saja (misal: hanya halaman Laba Rugi).")
        else:
            st.error(f"⚠️ Error: {err_msg}")

    st.rerun()

# ─────────────────────────────────────────────
# SAVE DATA
# ─────────────────────────────────────────────
if user:
    sessions_to_save = []
    for s in st.session_state.sessions:
        msgs = [dict(m) for m in s["messages"] if m["role"] != "system"]
        sessions_to_save.append({"id": s["id"], "title": s["title"], "created": s["created"], "messages": msgs})
    save_user(user["email"], {
        "theme": st.session_state.get("theme", "dark"),
        "sessions": sessions_to_save,
        "active_id": st.session_state.active_id,
    })

# Kirim token baru ke localStorage setelah login
_new_token = st.session_state.pop("new_token", None)
if _new_token:
    components.html(f"""
<script>
try {{ localStorage.setItem('sigma_token', '{_new_token}'); }} catch(e) {{}}
</script>
""", height=0)

# Auto-restore token saat refresh — baca localStorage lalu redirect
if st.session_state.user is None:
    components.html("""
<script>
(function() {
    try {
        var token = localStorage.getItem('sigma_token');
        if (token) {
            var url = window.parent.location.href.split('?')[0];
            window.parent.location.replace(url + '?sigma_token=' + token);
        }
    } catch(e) {}
})();
</script>
""", height=0)

# ─────────────────────────────────────────────
# JS: bubble kanan + paste gambar + mobile CSS
# ─────────────────────────────────────────────
components.html(f"""
<script>
const BC = "{C['bubble']}";
const BT = "#ffffff";

// ── Inject mobile CSS ke parent <head> — cara paling kuat ──
(function() {{
    var pd = window.parent.document;
    if (pd.getElementById('sigma-mobile-css')) return;
    var s = pd.createElement('style');
    s.id = 'sigma-mobile-css';
    s.textContent = `
        /* Base font konsisten */
        [data-testid="stMarkdownContainer"] p,
        [data-testid="stMarkdownContainer"] li,
        [data-testid="stMarkdownContainer"] span,
        [data-testid="stMarkdownContainer"] div,
        [data-testid="stMarkdownContainer"] strong,
        [data-testid="stMarkdownContainer"] b,
        [data-testid="stMarkdownContainer"] em {{
            font-size: 1rem !important;
            line-height: 1.85 !important;
        }}

        @media (max-width: 768px) {{
            /* Konten full width rapat */
            [data-testid="stMainBlockContainer"] {{
                max-width: 100% !important;
                padding: 8px 12px 120px !important;
                margin: 0 !important;
            }}
            /* Semua teks AI besar dan nyaman */
            [data-testid="stMarkdownContainer"],
            [data-testid="stMarkdownContainer"] p,
            [data-testid="stMarkdownContainer"] li,
            [data-testid="stMarkdownContainer"] span,
            [data-testid="stMarkdownContainer"] div,
            [data-testid="stMarkdownContainer"] strong,
            [data-testid="stMarkdownContainer"] b,
            [data-testid="stMarkdownContainer"] em,
            [data-testid="stMarkdownContainer"] a {{
                font-size: 1.05rem !important;
                line-height: 1.9 !important;
            }}
            [data-testid="stMarkdownContainer"] h1 {{ font-size: 1.3rem !important; }}
            [data-testid="stMarkdownContainer"] h2 {{ font-size: 1.15rem !important; }}
            [data-testid="stMarkdownContainer"] h3 {{ font-size: 1.05rem !important; font-weight: 700 !important; }}

            /* List rapi */
            [data-testid="stMarkdownContainer"] ul,
            [data-testid="stMarkdownContainer"] ol {{
                padding-left: 18px !important;
                margin: 4px 0 !important;
            }}
            [data-testid="stMarkdownContainer"] li {{
                margin-bottom: 6px !important;
            }}

            /* Jarak antar chat */
            [data-testid="stChatMessage"] {{
                padding: 10px 0 !important;
            }}

            /* Chat input */
            div[data-testid="stChatInputContainer"] {{
                border-radius: 26px !important;
                margin: 0 4px 8px !important;
            }}
            [data-testid="stChatInput"] textarea {{
                font-size: 16px !important;
                line-height: 1.5 !important;
            }}

            /* Bubble */
            .navy-pill {{
                max-width: 82% !important;
                font-size: 1.05rem !important;
                line-height: 1.75 !important;
                padding: 12px 16px !important;
            }}

            /* Code */
            [data-testid="stMarkdownContainer"] code {{
                font-size: 0.88rem !important;
            }}
            [data-testid="stMarkdownContainer"] pre {{
                font-size: 0.85rem !important;
                overflow-x: auto !important;
                padding: 12px !important;
            }}
        }}
    `;
    pd.head.appendChild(s);
}})();

function fixBubbles() {{
    const doc = window.parent.document;
    doc.querySelectorAll('[data-testid="stChatMessage"]').forEach(msg => {{
        if (!msg.querySelector('[data-testid="stChatMessageAvatarUser"]')) return;
        msg.style.cssText += 'display:flex!important;justify-content:flex-end!important;background:transparent!important;border:none!important;box-shadow:none!important;padding:4px 0!important;';
        const av = msg.querySelector('[data-testid="stChatMessageAvatarUser"]');
        if (av) av.style.display = 'none';
        const ct = msg.querySelector('[data-testid="stChatMessageContent"]');
        if (ct) ct.style.cssText += 'background:transparent!important;display:flex!important;justify-content:flex-end!important;max-width:100%!important;padding:0!important;';
        msg.querySelectorAll('[data-testid="stMarkdownContainer"]').forEach(md => {{
            md.style.background = 'transparent';
            md.style.display = 'flex';
            md.style.justifyContent = 'flex-end';
            if (!md.querySelector('.navy-pill')) {{
                const pill = document.createElement('div');
                pill.className = 'navy-pill';
                var mob=window.parent.innerWidth<=768;
                pill.style.cssText=`background:linear-gradient(135deg,#42a8e0,#1a4fad);color:#ffffff;border-radius:18px 18px 4px 18px;padding:${{mob?"12px 16px":"10px 16px"}};max-width:${{mob?"85%":"72%"}};display:inline-block;font-size:${{mob?"1rem":"0.9rem"}};line-height:1.7;word-wrap:break-word;`;
                while (md.firstChild) pill.appendChild(md.firstChild);
                md.appendChild(pill);
            }}
            // Force putih semua elemen dalam bubble
            var pill = md.querySelector('.navy-pill');
            if (pill) {{
                pill.style.setProperty('color','#ffffff','important');
                pill.style.setProperty('background','linear-gradient(135deg,#42a8e0,#1a4fad)','important');
                pill.querySelectorAll('*').forEach(function(el){{
                    el.style.setProperty('color','#ffffff','important');
                }});
            }}
        }});
    }});
}}

fixBubbles();
setInterval(fixBubbles, 800);
new MutationObserver(() => setTimeout(fixBubbles, 100)).observe(
    window.parent.document.body, {{childList:true,subtree:true}}
);

// ── Tombol aksi di bawah bubble & pesan AI ──
function addActionButtons() {{
    var doc = window.parent.document;
    doc.querySelectorAll('[data-testid="stChatMessage"]').forEach(function(msg) {{
        if (msg.querySelector('.sigma-actions')) return;
        if (!!msg.querySelector('[data-testid="stChatMessageAvatarUser"]')) return;
        function getMsgText() {{
            var md = msg.querySelector('[data-testid="stMarkdownContainer"]');
            return md ? md.innerText : '';
        }}
        var bar = doc.createElement('div');
        bar.className = 'sigma-actions';
        bar.style.cssText = 'display:flex;gap:2px;margin-top:6px;padding:0 2px;';
        var copyBtn = doc.createElement('button');
        copyBtn.title = 'Salin';
        copyBtn.innerHTML = '<svg width=\"16\" height=\"16\" viewBox=\"0 0 24 24\" fill=\"none\" stroke=\"#8e8ea0\" stroke-width=\"2\" stroke-linecap=\"round\" stroke-linejoin=\"round\"><rect x=\"9\" y=\"9\" width=\"13\" height=\"13\" rx=\"2\"></rect><path d=\"M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1\"></path></svg>';
        copyBtn.style.cssText = 'background:transparent;border:none;cursor:pointer;padding:5px 6px;border-radius:6px;display:flex;align-items:center;';
        copyBtn.onmouseenter=function(){{this.style.background='rgba(255,255,255,0.08)'}};
        copyBtn.onmouseleave=function(){{this.style.background='transparent'}};
        copyBtn.onclick = function() {{
            var txt = getMsgText();
            function showOk() {{
                copyBtn.innerHTML = '<svg width=\"16\" height=\"16\" viewBox=\"0 0 24 24\" fill=\"none\" stroke=\"#4CAF50\" stroke-width=\"2.5\" stroke-linecap=\"round\" stroke-linejoin=\"round\"><polyline points=\"20 6 9 17 4 12\"></polyline></svg>';
                setTimeout(function(){{ copyBtn.innerHTML = '<svg width=\"16\" height=\"16\" viewBox=\"0 0 24 24\" fill=\"none\" stroke=\"#8e8ea0\" stroke-width=\"2\" stroke-linecap=\"round\" stroke-linejoin=\"round\"><rect x=\"9\" y=\"9\" width=\"13\" height=\"13\" rx=\"2\"></rect><path d=\"M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1\"></path></svg>'; }}, 2000);
            }}
            navigator.clipboard.writeText(txt).then(showOk).catch(function(){{
                var ta=doc.createElement('textarea'); ta.value=txt;
                doc.body.appendChild(ta); ta.select(); doc.execCommand('copy'); doc.body.removeChild(ta);
                showOk();
            }});
        }};
        bar.appendChild(copyBtn);
        msg.style.flexDirection='column';
        msg.appendChild(bar);
    }});
}}

setInterval(addActionButtons, 1000);

// Paste image support — lebih robust
function setupPaste() {{
    var pw = window.parent;

    // Hapus handler lama dulu
    if (pw._sigmaPasteHandler) {{
        pw.removeEventListener('paste', pw._sigmaPasteHandler, true);
        pw.document.removeEventListener('paste', pw._sigmaPasteHandler, true);
    }}

    function handlePaste(e) {{
        var items = e.clipboardData && e.clipboardData.items;
        if (!items) return;
        for (var i=0; i<items.length; i++) {{
            if (items[i].type.startsWith('image/')) {{
                var file = items[i].getAsFile();
                if (!file) continue;
                e.preventDefault();
                e.stopPropagation();
                var inputs = pw.document.querySelectorAll('input[type="file"]');
                for (var fi of inputs) {{
                    try {{
                        var dt = new DataTransfer();
                        dt.items.add(file);
                        Object.defineProperty(fi, 'files', {{value: dt.files, configurable:true, writable:true}});
                        fi.dispatchEvent(new Event('change', {{bubbles:true}}));
                        fi.dispatchEvent(new Event('input', {{bubbles:true}}));
                        var ta = pw.document.querySelector('[data-testid="stChatInput"] textarea');
                        if (ta) {{
                            ta.style.outline = '2px solid #4a90d9';
                            ta.placeholder = '📎 Gambar siap — ketik pertanyaan lalu Enter';
                            setTimeout(function(){{
                                ta.style.outline='';
                                ta.placeholder='Tanya SIGMA... DYOR - bukan financial advice.';
                            }}, 3000);
                            ta.focus();
                        }}
                        break;
                    }} catch(err) {{ console.log('paste err',err); }}
                }}
                break;
            }}
        }}
    }}

    pw._sigmaPasteHandler = handlePaste;
    pw.addEventListener('paste', handlePaste, true);
    pw.document.addEventListener('paste', handlePaste, true);
}}
setupPaste();
setTimeout(setupPaste, 1000);
setTimeout(setupPaste, 3000);

// ── Drag & Drop file (PDF/gambar) ke area chat ──
function setupDragDrop() {{
    var pw = window.parent;
    var pd = pw.document;
    if (pw._sigmaDragOK) return;

    // Overlay saat drag
    var overlay = pd.createElement('div');
    overlay.id = 'sigma-drop-overlay';
    overlay.style.cssText = 'position:fixed;inset:0;background:rgba(27,42,74,0.55);z-index:99997;display:none;align-items:center;justify-content:center;pointer-events:none;';
    overlay.innerHTML = '<div style="background:#1B2A4A;color:#fff;border:2px dashed #4a90d9;border-radius:16px;padding:32px 48px;font-size:1.1rem;text-align:center;">📂 Lepaskan file di sini<br><span style="font-size:0.85rem;opacity:0.7;">PDF, PNG, JPG</span></div>';
    pd.body.appendChild(overlay);

    var dragCount = 0;

    pd.addEventListener('dragenter', function(e) {{
        e.preventDefault();
        dragCount++;
        overlay.style.display = 'flex';
    }}, true);

    pd.addEventListener('dragleave', function(e) {{
        dragCount--;
        if (dragCount <= 0) {{ dragCount = 0; overlay.style.display = 'none'; }}
    }}, true);

    pd.addEventListener('dragover', function(e) {{
        e.preventDefault();
    }}, true);

    pd.addEventListener('drop', function(e) {{
        e.preventDefault();
        dragCount = 0;
        overlay.style.display = 'none';

        var files = e.dataTransfer && e.dataTransfer.files;
        if (!files || files.length === 0) return;

        var file = files[0];
        var allowed = ['application/pdf','image/png','image/jpeg','image/jpg'];
        if (!allowed.includes(file.type)) {{
            alert('File tidak didukung. Gunakan PDF, PNG, atau JPG.');
            return;
        }}

        // Inject ke file input Streamlit
        var inputs = pd.querySelectorAll('input[type="file"]');
        for (var fi of inputs) {{
            try {{
                var dt = new DataTransfer();
                dt.items.add(file);
                Object.defineProperty(fi, 'files', {{value: dt.files, configurable:true, writable:true}});
                fi.dispatchEvent(new Event('change', {{bubbles:true}}));
                fi.dispatchEvent(new Event('input', {{bubbles:true}}));
                var ta = pd.querySelector('[data-testid="stChatInput"] textarea');
                if (ta) {{
                    ta.style.outline = '2px solid #4a90d9';
                    var fname = file.name;
                    ta.placeholder = '📎 ' + fname + ' siap — ketik pertanyaan lalu Enter';
                    setTimeout(function(){{
                        ta.style.outline = '';
                        ta.placeholder = 'Tanya SIGMA... DYOR - bukan financial advice.';
                    }}, 3000);
                    ta.focus();
                }}
                break;
            }} catch(err) {{ console.log('drop err', err); }}
        }}
    }}, true);

    pw._sigmaDragOK = true;
}}
setupDragDrop();
setTimeout(setupDragDrop, 2000);
</script>
""", height=0)
